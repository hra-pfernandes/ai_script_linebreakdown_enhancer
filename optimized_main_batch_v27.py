# Claude SQL Batch INSERT Pipeline  v11
# ─────────────────────────────────────────────────────────────────────────────
# Strategy   : Pure additive insertions — no replacement of existing SQL.
#
# Four REQUIRED insertion phases (minimum — more allowed per script):
#   1. cleanup_drop             → TOP of script, cleanup block
#   2. breakdown_block          → after HIERARCHY, before TOTAL_PRICE
#   3. final_join_detail_column → inside FINAL_JOIN after ExpectedDetailed
#   4. auditor_queries          → END of script, after final OUTPUT SELECT
#
# Anchoring  : Discovery derives anchors from FUNCTIONAL ROLE identification
#              (Phase 1 of system prompt), not arbitrary text snippets.
#
# Line nums  : anchor_line_number and insertion_ends_at_line are captured for
#              every insertion — making patched SQL easy to navigate/review.
#
# DB logging : position column is NVARCHAR(500) — stores anchor text.
#              anchor_line_number and insertion_ends_at_line are INT columns.
#              Run the migration script below before first use:
#
#   ALTER TABLE dbo.SQL_Processing_Edits
#       ALTER COLUMN position NVARCHAR(500);
#   ALTER TABLE dbo.SQL_Processing_Edits
#       ADD anchor_line_number    INT NULL,
#           insertion_ends_at_line INT NULL;
# ─────────────────────────────────────────────────────────────────────────────

import json
import os
import re
import random
import time
from datetime import datetime
from pathlib import Path

import anthropic
from database_manager import DatabaseManager
from docx import Document
from dotenv import load_dotenv

load_dotenv()

MODEL = "claude-sonnet-4-6"

INPUT_COST_PER_MILLION      = 3.00
OUTPUT_COST_PER_MILLION     = 15.00
CACHE_READ_COST_PER_MILLION = 0.15

BASE_DIR           = Path(__file__).resolve().parent
SYSTEM_PROMPT_PATH = BASE_DIR / "data" / "HRA_System_Prompt_v16.docx"
REFERENCE_SQL_PATH = BASE_DIR / "data" / "NYP_COL_GHI_CBP_Commercial_OP_Complete.sql"
INPUT_DIR          = BASE_DIR / "input"
BATCH_OUTPUT_DIR   = BASE_DIR / "batch_output"

CONTEXT_CHARS_BEFORE = 1000
CONTEXT_CHARS_AFTER  = 1000

# Every script must have at least these four phases.
REQUIRED_PHASES = {
    "cleanup_drop",
    "breakdown_block",
    "final_join_detail_column",
    "auditor_queries",
}

total_cached_tokens   = 0
total_uncached_tokens = 0

session_metrics = {
    "start_time": None,
    "files_processed": 0,
    "files_successful": 0,
    "files_failed": 0,
    "total_insertions": 0,
    "total_successful_insertions": 0,
    "total_cost": 0.0,
    "total_savings": 0.0,
}

db_manager = None


# ─────────────────────────────────────────────────────────────────────────────
# I/O helpers
# ─────────────────────────────────────────────────────────────────────────────

def read_docx(path):
    doc   = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for r in t.rows:
            cells = [c.text.strip() for c in r.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_text(path):
    return path.read_text(encoding="utf-8", errors="ignore")


def append_log(log_path, line):
    log_path.parent.mkdir(parents=True, exist_ok=True)
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(line + "\n")


def create_file_output_structure(input_filename, batch_output_dir):
    base_name       = Path(input_filename).stem
    file_output_dir = batch_output_dir / base_name
    file_output_dir.mkdir(parents=True, exist_ok=True)
    return {
        "output_dir":     file_output_dir,
        "patched_sql":    file_output_dir / f"{base_name}_patched.sql",
        "log_file":       file_output_dir / f"{base_name}_metrics.log",
        "insertions_log": file_output_dir / f"{base_name}_insertions.json",
        "analysis_log":   file_output_dir / f"{base_name}_analysis.json",
    }


# ─────────────────────────────────────────────────────────────────────────────
# Cost helpers
# ─────────────────────────────────────────────────────────────────────────────

def estimate_cost(input_tokens, output_tokens, cached_tokens):
    return (
          (input_tokens  / 1_000_000) * INPUT_COST_PER_MILLION
        + (cached_tokens / 1_000_000) * CACHE_READ_COST_PER_MILLION
        + (output_tokens / 1_000_000) * OUTPUT_COST_PER_MILLION
    )


def usage_numbers(msg):
    u = getattr(msg, "usage", None)
    if not u:
        return 0, 0, 0, 0
    return (
        getattr(u, "input_tokens", 0),
        getattr(u, "output_tokens", 0),
        getattr(u, "cache_creation_input_tokens", 0),
        getattr(u, "cache_read_input_tokens", 0),
    )


# ─────────────────────────────────────────────────────────────────────────────
# JSON extraction
# ─────────────────────────────────────────────────────────────────────────────

def extract_json(text):
    """
    Extract and parse a JSON object from Claude's response.

    Handles three common response formats:
      1. Raw JSON   — Claude returns { ... } directly
      2. Code fence — Claude wraps in ```json ... ``` or ``` ... ```
      3. Mixed text — JSON embedded somewhere in prose

    Enhanced with control character cleaning for SQL code in JSON strings.
    """
    if not text or not text.strip():
        return {"error": "Empty or whitespace-only response", "raw_text": text}
    
    text = text.strip()
    original_length = len(text)

    def clean_json_for_parsing(json_text):
        """Clean control characters that break JSON parsing while preserving structure."""
        # Replace problematic control characters in string values
        # This regex finds string values and replaces control chars within them
        import re
        
        def replace_in_string(match):
            string_content = match.group(0)
            # Replace literal newlines, tabs, etc. with escaped versions
            cleaned = string_content.replace('\\n', '\\n')  # Already escaped - keep as is
            cleaned = cleaned.replace('\n', '\\n')          # Literal newline - escape it
            cleaned = cleaned.replace('\r', '\\r')          # Literal carriage return
            cleaned = cleaned.replace('\t', '\\t')          # Literal tab
            cleaned = cleaned.replace('\b', '\\b')          # Literal backspace
            cleaned = cleaned.replace('\f', '\\f')          # Literal form feed
            return cleaned
        
        # Find all string values (including multiline ones) and clean them
        # This regex matches: "..." including escaped quotes and newlines
        pattern = r'"(?:[^"\\]|\\.)*"'
        return re.sub(pattern, replace_in_string, json_text, flags=re.DOTALL)

    # Strategy 1: strip markdown code fences if present, then parse
    fence_stripped = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    fence_stripped = re.sub(r"\s*```\s*$", "", fence_stripped).strip()
    
    if fence_stripped != text:
        print(f"  Stripped code fences: {len(text)} -> {len(fence_stripped)} chars")
    
    # Clean control characters before parsing
    cleaned_text = clean_json_for_parsing(fence_stripped)
    if cleaned_text != fence_stripped:
        print(f"  Cleaned control characters: {len(fence_stripped)} -> {len(cleaned_text)} chars")
    
    try:
        parsed = json.loads(cleaned_text)
        print(f"  JSON parsed successfully (fence-stripped + cleaned, {len(cleaned_text)} chars)")
        return parsed
    except json.JSONDecodeError as e:
        print(f"  Fence-stripped + cleaned parse failed: {e}")

    # Strategy 2: direct parse (no fences, but still clean)
    cleaned_direct = clean_json_for_parsing(text)
    try:
        parsed = json.loads(cleaned_direct)
        print(f"  JSON parsed successfully (direct + cleaned, {len(cleaned_direct)} chars)")
        return parsed
    except json.JSONDecodeError as e:
        print(f"  Direct + cleaned parse failed: {e}")

    # Strategy 3: extract outermost { ... } using rfind for the closing brace
    start = text.find("{")
    end   = text.rfind("}")
    
    if start == -1:
        print(f"  No opening brace found in {len(text)} chars")
        return {"error": "No JSON opening brace found", "raw_text": text[:500]}
    
    if end == -1:
        print(f"  No closing brace found in {len(text)} chars")
        return {"error": "No JSON closing brace found - possibly truncated", "raw_text": text[:500]}
    
    if end <= start:
        print(f"  Invalid brace positions: start={start}, end={end}")
        return {"error": "Invalid JSON brace positions", "raw_text": text[:500]}
    
    json_candidate = text[start : end + 1]
    print(f"  Extracting JSON candidate: {len(json_candidate)} chars from positions {start}-{end}")
    
    # Clean the extracted candidate
    cleaned_candidate = clean_json_for_parsing(json_candidate)
    
    try:
        parsed = json.loads(cleaned_candidate)
        print(f"  JSON parsed successfully (extracted + cleaned, {len(cleaned_candidate)} chars)")
        return parsed
    except json.JSONDecodeError as e:
        print(f"  Extracted + cleaned JSON parse failed: {e}")
        print(f"  JSON candidate preview: {json_candidate[:200]}...")
        
        # Check for common truncation indicators
        if json_candidate.endswith("...") or "..." in json_candidate:
            return {"error": "JSON appears truncated (contains ellipsis)", "raw_text": text[:500]}

    # All strategies failed - provide detailed error info
    error_info = {
        "error": "Could not parse JSON with any strategy",
        "original_length": original_length,
        "cleaned_length": len(cleaned_text),
        "has_opening_brace": start != -1,
        "has_closing_brace": end != -1,
        "brace_positions": f"{start}-{end}" if start != -1 and end != -1 else "N/A",
        "contains_ellipsis": "..." in text,
        "raw_text": text[:500]
    }
    
    print(f"  JSON parse error. Length: {original_length}, Braces: {start}-{end}, Preview: {text[:100]}...")
    return error_info


# ─────────────────────────────────────────────────────────────────────────────
# Line number helper
# ─────────────────────────────────────────────────────────────────────────────

def get_line_number(sql: str, char_pos: int) -> int:
    """Return 1-based line number for a character position in sql."""
    return sql[:char_pos].count("\n") + 1


def validate_branked_consistency(new_content: str) -> list[str]:
    """
    After breakdown_block generation, verify that every rn_ alias referenced
    in #LineBreakdown is also defined in #bRanked within the same new_content.

    Also detects template placeholder aliases (e.g. rn_CAT, rn_ALIAS) that
    the model copied verbatim from the instruction examples instead of
    substituting the real category name.

    Returns a list of missing/placeholder alias names (empty = all good).
    """
    import re

    # Template placeholder names that should never appear in real SQL.
    PLACEHOLDERS = {
        "rn_cat", "rn_alias", "rn_max_alias", "rn_hybrid_alias",
        "rn_window_alias", "rn_category_alias", "rn_actualcategoryname",
        "rn_actualcategory",
    }

    # Strip single-line SQL comments (-- ...) before scanning.
    # This prevents rn_ aliases that appear only inside comment text from
    # being flagged as missing — the model often writes comment documentation
    # using example patterns like IIF(b.rn_CAT = 1, ...) which are not real
    # alias references even though they look like one to a regex scanner.
    def strip_sql_comments(sql: str) -> str:
        # Remove -- line comments but preserve the newline
        return re.sub(r'--[^\n]*', '', sql)

    code_only = strip_sql_comments(new_content)

    # Find all rn_ aliases DEFINED in #bRanked (AS rn_xxx pattern)
    # Scan full content including comments to catch definitions inside comment blocks
    defined    = set(re.findall(r'\bAS\s+(rn_\w+)', new_content, re.IGNORECASE))
    # Find all rn_ aliases REFERENCED in actual code (comments stripped)
    referenced = set(re.findall(r'\b(rn_\w+)\b', code_only, re.IGNORECASE))

    # Missing = referenced in code but not defined
    missing = [r for r in referenced if r not in defined]

    # Placeholders in actual code (not comments) = template names that survived
    placeholders = [a for a in (defined | referenced) if a.lower() in PLACEHOLDERS]

    return list(set(missing + placeholders))



def validate_breakdown_source(new_content: str, analysis: dict) -> list[str]:
    """
    After breakdown_block generation, verify that #bRanked is sourced from
    the correct line-level source table (e.g. #Step3), not an earlier table
    like #Step2.

    The LEAD slot columns (AS1-AS10, Rad1-Rad10) only exist in the
    line-level source table. Using #Step2 or any pre-LEAD table causes a
    guaranteed runtime crash on #bSlots.

    Also checks that #bSlots does NOT contain recomputed slot division logic
    (a sign the model tried to reconstruct slots from scratch rather than
    reading them from the source table).

    Returns a list of warning strings (empty = all good). These are FATAL.
    """
    import re as _re
    warnings = []
    source_table = analysis.get("line_level_source_table", "").strip().lower()

    # Find the FROM clause of the SELECT INTO #bRanked block
    branked_from = _re.search(
        r'INTO\s+#bRanked\s*[\r\n\s]+FROM\s+(#\w+)',
        new_content, _re.IGNORECASE
    )
    if branked_from:
        actual_source = branked_from.group(1).lower()
        if source_table and actual_source != source_table:
            warnings.append(
                f"WRONG_BRANKED_SOURCE | #bRanked is sourced from "
                f"'{branked_from.group(1)}' but the line-level source table "
                f"identified in Phase 1A is '{analysis.get('line_level_source_table')}'. "
                f"#bRanked MUST be sourced from the line-level source table because "
                f"it is the only table containing the LEAD slot columns (AS1-AS10, "
                f"Rad1-Rad10). Using an earlier table causes a runtime crash on #bSlots."
            )

    # Check for recomputed slot division logic in #bSlots
    # Pattern: ISNULL([AS], 0) / 2.0 or [Mammography_Reduction] / 2.0
    slot_recompute = _re.search(
        r'ISNULL\s*\(\s*\[?(AS|Mammography_Reduction)\]?\s*,\s*0\s*\)\s*/\s*[0-9]',
        new_content, _re.IGNORECASE
    )
    if slot_recompute:
        warnings.append(
            f"SLOT_RECOMPUTE_DETECTED | #bSlots contains slot recomputation "
            f"logic (dividing the raw category value by a divisor). This is "
            f"incorrect — slot values are already computed by the LEAD() "
            f"expressions in the line-level source table and should be read "
            f"directly using MAX(CASE WHEN rn_X=1 THEN SlotCol ELSE 0 END). "
            f"Recomputing slots applies the divisor to each line's own value "
            f"instead of distributing rank-1's value across all ranks."
        )

    return warnings


def validate_contract_years(new_content: str, analysis: dict) -> list[str]:
    """
    After breakdown_block generation, verify that ContractEffectiveDateFrom
    and ContractEffectiveDateTo CASE blocks are present in #LineBreakdown,
    and that period date boundaries from contract_periods all appear in the SQL.
    Returns a list of error strings (empty = all good). Errors are FATAL.
    """
    import re as _re
    warnings = []
    contract_periods = analysis.get("contract_periods", [])

    has_from = bool(_re.search(r'ContractEffectiveDateFrom', new_content, _re.IGNORECASE))
    has_to   = bool(_re.search(r'ContractEffectiveDateTo',   new_content, _re.IGNORECASE))

    if not has_from:
        warnings.append(
            "MISSING_CONTRACT_DATE_FROM | ContractEffectiveDateFrom column is absent "
            "from #LineBreakdown. A CASE block keyed on ServiceDate BETWEEN each "
            "contract period must be generated and aliased as ContractEffectiveDateFrom."
        )
    if not has_to:
        warnings.append(
            "MISSING_CONTRACT_DATE_TO | ContractEffectiveDateTo column is absent "
            "from #LineBreakdown. A CASE block keyed on ServiceDate BETWEEN each "
            "contract period must be generated and aliased as ContractEffectiveDateTo."
        )

    if has_from and contract_periods:
        def _date_found_in_sql(date_str: str, sql: str) -> bool:
            """
            Match a YYYY-MM-DD date against SQL that may use any of:
              'YYYY-MM-DD', 'MM/DD/YYYY', 'YYYYMMDD', 'YYYY/MM/DD'
            All wrapped in optional single-quotes.
            """
            if not date_str:
                return False
            # Fast path: exact ISO string present
            if date_str in sql:
                return True
            try:
                parts = date_str.split("-")
                if len(parts) != 3:
                    return False
                yyyy, mm, dd = parts
                # Build all common SQL date literal patterns
                patterns = [
                    rf"'?{yyyy}-{mm}-{dd}'?",      # ISO:      2023-01-01
                    rf"'?{mm}/{dd}/{yyyy}'?",        # US:       01/01/2023
                    rf"'?{yyyy}/{mm}/{dd}'?",        # Alt ISO:  2023/01/01
                    rf"'?{yyyy}{mm}{dd}'?",          # Compact:  20230101
                ]
                return any(_re.search(p, sql) for p in patterns)
            except Exception:
                return date_str in sql  # fallback to exact match

        # Deduplicate by date_from — multiple periods can share the same start date
        # (e.g. a Drugs period and a main period both starting 2022-04-01).
        # The CASE block only needs each unique date_from once; counting raw
        # len(contract_periods) against a set of unique dates always fails when
        # duplicates exist.
        unique_dates_from = sorted({
            p["date_from"] for p in contract_periods
            if isinstance(p, dict) and p.get("date_from")
        })
        expected_count = len(unique_dates_from)

        found_dates = {
            d for d in unique_dates_from
            if _date_found_in_sql(d, new_content)
        }
        if len(found_dates) < expected_count:
            missing_dates = [d for d in unique_dates_from if d not in found_dates]
            warnings.append(
                f"CONTRACT_PERIOD_GAP | ContractEffectiveDateFrom covers "
                f"{len(found_dates)} of {expected_count} unique date_from values. "
                f"Missing: {missing_dates}. "
                f"Dates must appear as ISO literals (YYYY-MM-DD) in the CASE WHEN BETWEEN blocks."
            )
    return warnings


def validate_hybrid_classification(analysis: dict) -> list[str]:
    """
    Cross-check all category classifications against the raw aggregation
    evidence recorded in aggregate_date_raw and aggregate_enc_raw.

    Returns a list of warning strings (empty = all good).
    Warnings are non-fatal — logged and printed for operator review before
    breakdown_block generation proceeds.

    Rules enforced — each warning includes a severity tag:

    FATAL (blocks breakdown_block generation, triggers Discovery retry):
      HYBRID_MISS   : MAX in agg_date + SUM in agg_enc but not in hybrid_categories
      HYBRID_WRONG  : in hybrid_categories but evidence is not MAX/SUM
      SUM_IS_HYBRID : in sum_categories but evidence is MAX(agg_date)/SUM(agg_enc)
      SUM_IS_MAX    : in sum_categories but agg_enc shows MAX
      MAX_IS_SUM    : in max_categories but agg_enc shows SUM
      INDICATOR_IN_MAX : INDICATOR_FLAG category also listed in max_categories

    WARNING (logged, printed, non-blocking):
      MAX_IS_HYBRID : in max_categories but agg_date=SUM, agg_enc=MAX (unusual pattern)
      LEAD_PARTITION_MISMATCH  : branked_partition_by != lead_partition_by
      BSLOTS_PARTITION_MISMATCH: bslots_group_by != lead_partition_by
    """
    FATAL_CODES = {
        "HYBRID_MISS", "HYBRID_WRONG", "SUM_IS_HYBRID",
        "SUM_IS_MAX", "MAX_IS_SUM", "INDICATOR_IN_MAX",
        "WINDOW_REDUCTION_IS_HYBRID",
    }

    issues = []   # list of dicts: {code, fatal, message}

    def add(code, msg):
        issues.append({
            "code":    code,
            "fatal":   code in FATAL_CODES,
            "message": msg,
        })

    has_agg_date = analysis.get("has_aggregate_date_step", False)

    agg_date_raw = {
        row["column"].lower(): row["agg_fn"].upper()
        for row in analysis.get("aggregate_date_raw", {}).get("columns", [])
        if isinstance(row, dict) and isinstance(row.get("column"), str)
        and isinstance(row.get("agg_fn"), str)
    }
    agg_enc_raw = {
        row["column"].lower(): row["agg_fn"].upper()
        for row in analysis.get("aggregate_enc_raw", {}).get("columns", [])
        if isinstance(row, dict) and isinstance(row.get("column"), str)
        and isinstance(row.get("agg_fn"), str)
    }

    declared_hybrid   = {c.lower() for c in analysis.get("hybrid_categories", []) if isinstance(c, str)}
    declared_max      = {c.lower() for c in analysis.get("max_categories", []) if isinstance(c, str)}
    declared_sum      = {c.lower() for c in analysis.get("sum_categories", []) if isinstance(c, str)}
    declared_flags    = {c.lower() for c in analysis.get("indicator_flag_categories", []) if isinstance(c, str)}

    # WINDOW_REDUCTION categories are exempt from aggregation-based MAX/HYBRID checks.
    # Their correct classification is determined by their LEAD partition scope, not
    # by the aggregation evidence. patch_analysis_classifications enforces this.
    window_reduction_cats = {
        e.get("category", "").lower()
        for e in analysis.get("window_reduction_lead_partitions", [])
        if e.get("category")
    }

    # ── INDICATOR_IN_MAX: flag also listed in max_categories ─────────────────────
    for col in declared_flags & declared_max:
        add("INDICATOR_IN_MAX",
            f"'{col}' is in both indicator_flag_categories and max_categories. "
            f"INDICATOR_FLAG categories must not appear in max_categories.")

    # ── HYBRID detection ───────────────────────────────────────────────────────────────
    if agg_date_raw:
        for col, date_fn in agg_date_raw.items():
            enc_fn = agg_enc_raw.get(col)
            if date_fn == "MAX" and enc_fn == "SUM":
                if col in window_reduction_cats:
                    continue  # WINDOW_REDUCTION: classification driven by LEAD, not agg
                if col not in declared_hybrid:
                    bucket = (
                        "sum_categories" if col in declared_sum else
                        "max_categories" if col in declared_max else
                        "unclassified"
                    )
                    add("HYBRID_MISS",
                        f"'{col}' is MAX in AGGREGATE_DATE but SUM in AGGREGATE_ENC "
                        f"— must be in hybrid_categories. Currently declared as: {bucket}. "
                        f"If left as SUM, LinePayment pays every line instead of one winner "
                        f"per date, causing overpayment invisible to reconciliation.")

    for col in declared_hybrid:
        if col in window_reduction_cats:
            continue  # WINDOW_REDUCTION: LEAD scope determines classification
        date_fn = agg_date_raw.get(col)
        enc_fn  = agg_enc_raw.get(col)
        if date_fn and enc_fn:
            if not (date_fn == "MAX" and enc_fn == "SUM"):
                add("HYBRID_WRONG",
                    f"'{col}' is in hybrid_categories but evidence is "
                    f"AGGREGATE_DATE={date_fn}, AGGREGATE_ENC={enc_fn}. "
                    f"Expected AGGREGATE_DATE=MAX, AGGREGATE_ENC=SUM.")

    # ── MAX validation ─────────────────────────────────────────────────────────────────────
    for col in declared_max - declared_flags:
        if col in window_reduction_cats:
            continue  # WINDOW_REDUCTION: classification driven by LEAD scope, not agg evidence
        date_fn = agg_date_raw.get(col)
        enc_fn  = agg_enc_raw.get(col)
        if has_agg_date and date_fn == "SUM" and enc_fn == "MAX":
            add("MAX_IS_HYBRID",   # warning only
                f"'{col}' is in max_categories but evidence is "
                f"AGGREGATE_DATE=SUM, AGGREGATE_ENC=MAX — unusual pattern. Verify.")
        if enc_fn == "SUM":
            add("MAX_IS_SUM",
                f"'{col}' is in max_categories but AGGREGATE_ENC uses SUM(). "
                f"Generated IIF(rn=1) logic will cause underpayment.")

    # ── SUM validation ─────────────────────────────────────────────────────────────────────
    for col in declared_sum:
        date_fn = agg_date_raw.get(col)
        enc_fn  = agg_enc_raw.get(col)
        if has_agg_date and date_fn == "MAX" and enc_fn == "SUM":
            add("SUM_IS_HYBRID",
                f"'{col}' is in sum_categories but evidence is "
                f"AGGREGATE_DATE=MAX, AGGREGATE_ENC=SUM — this is HYBRID. "
                f"Plain SUM LinePayment pays every line instead of one winner per date. "
                f"Overpayment WILL NOT be caught by reconciliation.")
        if enc_fn == "MAX":
            add("SUM_IS_MAX",
                f"'{col}' is in sum_categories but AGGREGATE_ENC uses MAX(). "
                f"Plain SUM LinePayment will overpay.")

    # ── WINDOW_REDUCTION checks ───────────────────────────────────────────────────────────────────
    for entry in analysis.get("window_reduction_lead_partitions", []):
        cat          = entry.get("category", "?")
        lead_part    = entry.get("lead_partition_by", "")
        branked_part = entry.get("branked_partition_by", "")
        bslots_grp   = entry.get("bslots_group_by", "")

        # WINDOW_REDUCTION_IS_HYBRID (FATAL): a WINDOW_REDUCTION category whose
        # LEAD partitions by EncounterID only must NOT be in hybrid_categories.
        # If the LEAD uses EncounterID only, #bRanked must use PARTITION BY
        # [EncounterID] only. Classifying it as HYBRID and using
        # PARTITION BY [EncounterID],[ServiceDate] produces mismatched rankings
        # because the slot values were pre-computed per-encounter, not per-date.
        # The aggregation-derived MAX→SUM pattern is misleading for these
        # categories — the LEAD partition scope is the authoritative rule.
        if lead_part == "EncounterID_only" and cat.lower() in declared_hybrid:
            add("WINDOW_REDUCTION_IS_HYBRID",
                f"'{cat}' is a WINDOW_REDUCTION category with lead_partition_by= "
                f"'{lead_part}' but is in hybrid_categories. WINDOW_REDUCTION "
                f"categories with EncounterID-only LEAD must be in max_categories "
                f"(PARTITION BY [EncounterID] in #bRanked). Using HYBRID partition "
                f"produces mismatched slot rankings and wrong LinePayment values.")

        # Partition consistency warnings (non-fatal)
        if lead_part and branked_part and lead_part != branked_part:
            add("LEAD_PARTITION_MISMATCH",
                f"'{cat}': LEAD uses '{lead_part}' but branked_partition_by='{branked_part}'.")
        if lead_part and bslots_grp and lead_part != bslots_grp:
            add("BSLOTS_PARTITION_MISMATCH",
                f"'{cat}': LEAD uses '{lead_part}' but bslots_group_by='{bslots_grp}'.")

    return issues, agg_date_raw, agg_enc_raw


# ─────────────────────────────────────────────────────────────────────────────
# Anchor helpers
# ─────────────────────────────────────────────────────────────────────────────


def patch_analysis_classifications(analysis: dict,
                                   agg_date_raw: dict,
                                   agg_enc_raw: dict) -> dict:
    """
    Directly rewrite the classification buckets in `analysis` using the raw
    aggregation evidence, with no LLM involvement.

    This is called on the FIRST retry instead of sending another API call.
    Since the model already returned correct evidence (agg_date_raw /
    agg_enc_raw) but put categories in the wrong buckets, we can derive the
    correct classification deterministically in Python.

    Rules applied (in priority order):
      1. INDICATOR_FLAG           → indicator_flag_categories only, never elsewhere
      2. WINDOW_REDUCTION with
         lead_partition_by=EncounterID_only
                                  → max_categories (one winner per encounter,
                                    not HYBRID — the LEAD already encodes
                                    across-encounter ordering, not per-date)
      3. WINDOW_REDUCTION with
         lead_partition_by=EncounterID_and_ServiceDate
                                  → hybrid_categories (one winner per date)
      4. agg_date=MAX, agg_enc=MAX → max_categories
      5. agg_date=SUM, agg_enc=SUM → sum_categories
      6. agg_date=MAX, agg_enc=SUM → hybrid_categories
      7. agg_date=SUM, agg_enc=MAX → max_categories  (unusual pattern)
      8. absent from enc_raw       → keep original bucket

    Rule 2/3 exist because WINDOW_REDUCTION categories computed via LEAD()
    have their slot values pre-calculated at the line level using the LEAD
    partition scope. The aggregation-derived HYBRID classification (MAX→SUM)
    reflects how the slot sums are aggregated across dates, not how individual
    line payments are determined. Using PARTITION BY EncounterID+ServiceDate
    in #bRanked for a LEAD that was computed PARTITION BY EncounterID produces
    mismatched rankings and wrong slot assignments.
    """
    import copy
    patched = copy.deepcopy(analysis)

    indicator_flags = {c.lower() for c in patched.get("indicator_flag_categories", []) if isinstance(c, str)}

    # Build a lookup of WINDOW_REDUCTION categories and their LEAD partition scope
    # from window_reduction_lead_partitions (set by the model in Phase 1E).
    window_reduction_partitions: dict[str, str] = {}
    for entry in patched.get("window_reduction_lead_partitions", []):
        cat  = entry.get("category", "").lower()
        part = entry.get("lead_partition_by", "EncounterID_only")
        if cat:
            window_reduction_partitions[cat] = part

    new_max    = []
    new_hybrid = []
    new_sum    = []

    # Collect all known category names from all buckets + raw evidence
    all_cats = set()
    for bucket in ("max_categories", "hybrid_categories", "sum_categories",
                   "indicator_flag_categories"):
        for c in patched.get(bucket, []):
            if isinstance(c, str):
                all_cats.add(c.lower())
    for c in list(agg_enc_raw.keys()) + list(agg_date_raw.keys()):
        all_cats.add(c.lower())

    for cat in all_cats:
        # Rule 1: INDICATOR_FLAGs stay in their own bucket — never touch them
        if cat in indicator_flags:
            continue

        # Rules 2 & 3: WINDOW_REDUCTION — follow LEAD partition, ignore agg evidence
        if cat in window_reduction_partitions:
            part = window_reduction_partitions[cat]
            if part == "EncounterID_and_ServiceDate":
                # LEAD partitioned by date → one winner per date → HYBRID
                new_hybrid.append(cat)
            else:
                # LEAD partitioned by encounter only → one winner per encounter → MAX
                # #bRanked will use PARTITION BY [EncounterID] only for this category.
                new_max.append(cat)
            continue

        date_fn = agg_date_raw.get(cat, "")
        enc_fn  = agg_enc_raw.get(cat, "")

        if not enc_fn:
            # Rule 8: No AGGREGATE_ENC evidence — keep original bucket
            for bucket in ("max_categories", "hybrid_categories", "sum_categories"):
                if any(c.lower() == cat for c in patched.get(bucket, [])):
                    if bucket == "max_categories":
                        new_max.append(cat)
                    elif bucket == "hybrid_categories":
                        new_hybrid.append(cat)
                    else:
                        new_sum.append(cat)
            continue

        # Rules 4-7: derive from aggregation evidence
        if date_fn == "MAX" and enc_fn == "SUM":
            new_hybrid.append(cat)   # Rule 6
        elif enc_fn == "MAX":
            new_max.append(cat)      # Rules 4 & 7
        else:
            new_sum.append(cat)      # Rule 5

    # Preserve original casing from whichever bucket the category was in
    def restore_case(names_lower: list, orig_analysis: dict) -> list:
        all_originals = {}
        for bucket in ("max_categories", "hybrid_categories", "sum_categories",
                       "indicator_flag_categories"):
            for c in orig_analysis.get(bucket, []):
                if isinstance(c, str):
                    all_originals[c.lower()] = c
        for c in list(agg_enc_raw.keys()) + list(agg_date_raw.keys()):
            if c.lower() not in all_originals:
                all_originals[c.lower()] = c
        return [all_originals.get(n, n) for n in names_lower]

    patched["max_categories"]    = restore_case(new_max,    analysis)
    patched["hybrid_categories"] = restore_case(new_hybrid, analysis)
    patched["sum_categories"]    = restore_case(new_sum,    analysis)

    return patched


def find_anchor(sql: str, anchor: str) -> int | None:
    """
    Return the character position immediately AFTER `anchor` in `sql`.
    Tries exact match first, then whitespace-normalised match.
    """
    # 1. Exact
    pos = sql.find(anchor)
    if pos != -1:
        return pos + len(anchor)

    # 2. Normalised
    sql_norm    = re.sub(r"\s+", " ", sql)
    anchor_norm = re.sub(r"\s+", " ", anchor.strip())
    pos_norm    = sql_norm.find(anchor_norm)
    if pos_norm != -1:
        target_non_ws = len(re.sub(r"\s+", "", sql_norm[: pos_norm + len(anchor_norm)]))
        seen = 0
        for i, ch in enumerate(sql):
            if not ch.isspace():
                seen += 1
            if seen >= target_non_ws:
                return i + 1

    return None


def validate_anchor_against_table(sql: str, anchor: str, table_name: str) -> bool:
    """
    Confirm table_name appears anywhere in the SQL preceding the anchor.
    Searches the full preceding text — a fixed window would fail after large
    insertions shift the anchor far from its role table's definition.
    """
    anchor_pos = find_anchor(sql, anchor)
    if anchor_pos is None:
        return False
    preceding = sql[:anchor_pos]
    return table_name.upper() in preceding.upper()


def get_context(sql: str, anchor: str) -> str:
    pos = find_anchor(sql, anchor)
    if pos is None:
        return sql[:3000]
    return sql[max(0, pos - CONTEXT_CHARS_BEFORE) : min(len(sql), pos + CONTEXT_CHARS_AFTER)]


def apply_insertion(sql: str, anchor: str, new_content: str,
                    position: str) -> tuple[str, bool]:
    """Insert new_content immediately after (or before) the anchor."""
    anchor_end = find_anchor(sql, anchor)
    if anchor_end is None:
        print(f"  ERROR: anchor not found — '{anchor[:80]}'")
        return sql, False
    insert_at = anchor_end if position == "after" else max(0, anchor_end - len(anchor))
    return sql[:insert_at] + "\n\n" + new_content.strip() + "\n\n" + sql[insert_at:], True


# ─────────────────────────────────────────────────────────────────────────────
# Discovery prompt
# ─────────────────────────────────────────────────────────────────────────────

# ── Call 1: analysis-only schema (no insertion planning) ────────────────────
# Focused entirely on classification. Smaller output, less competing pressure.
ANALYSIS_SCHEMA = """{
  "line_level_source_table":     "#ActualName",
  "hierarchy_table":             "#ActualName",
  "total_price_table":           "#ActualName",
  "final_join_table":            "#ActualName",
  "output_table":                "#ActualName",
  "has_ncci_bundle":             false,
  "has_window_reduction":        false,
  "window_reduction_categories": [],
  "has_aggregate_date_step":     false,
  "hybrid_categories":           [],
  "max_categories":              [],
  "sum_categories":              [],
  "indicator_flag_categories":   [],
  "aggregate_date_raw": {
    "columns": [{"column": "CategoryName", "agg_fn": "MAX or SUM"}]
  },
  "aggregate_enc_raw": {
    "columns": [{"column": "CategoryName", "agg_fn": "MAX or SUM"}]
  },
  "window_reduction_lead_partitions": [
    {
      "category": "CategoryName",
      "slot_columns": ["Slot1", "Slot2"],
      "lead_partition_by":    "EncounterID_only OR EncounterID_and_ServiceDate",
      "branked_partition_by": "EncounterID_only OR EncounterID_and_ServiceDate",
      "bslots_group_by":      "EncounterID_only OR EncounterID_and_ServiceDate"
    }
  ],
  "contract_periods": [
    {
      "period_label": "Year1",
      "date_from":    "YYYY-MM-DD",
      "date_to":      "YYYY-MM-DD"
    }
  ]
}"""

# ── Call 2: insertion-planning schema (uses validated analysis as input) ─────
PLANNING_SCHEMA = """{
  "insertions": [
    {
      "insertion_id":     1,
      "phase":            "cleanup_drop",
      "description":      "DROP TABLE safety lines for #bRanked, #bSlots, #LineBreakdown in the cleanup block",
      "role_anchor_table": "cleanup_block",
      "anchor_snippet":   "exact last ~80 chars of the last existing DROP TABLE line in the cleanup block at the TOP of the script",
      "insert_position":  "after",
      "occurrence_index": 1
    },
    {
      "insertion_id":     2,
      "phase":            "breakdown_block",
      "description":      "#bRanked + optional #bSlots + SELECT INTO #LineBreakdown",
      "role_anchor_table": "#HierarchyTableActualName",
      "anchor_snippet":   "exact last ~80 chars of the final statement of the HIERARCHY step",
      "insert_position":  "after",
      "occurrence_index": 1
    },
    {
      "insertion_id":     3,
      "phase":            "final_join_detail_column",
      "description":      "Price_Breakdown_Detail correlated subquery as last column in FINAL_JOIN SELECT",
      "role_anchor_table": "#FinalJoinTableActualName",
      "anchor_snippet":   "exact text of the INTO #TableName line of the FINAL_JOIN SELECT",
      "insert_position":  "before",
      "occurrence_index": 1
    },
    {
      "insertion_id":     4,
      "phase":            "auditor_queries",
      "description":      "Auditor Query 1 + INSERT + Query 2 reconciliation appended after final OUTPUT SELECT",
      "role_anchor_table": "#OutputTableActualName",
      "anchor_snippet":   "exact last ~80 chars of the final OUTPUT SELECT statement",
      "insert_position":  "after",
      "occurrence_index": 1
    }
  ]
}"""

# Legacy alias so existing references to DISCOVERY_SCHEMA still compile
DISCOVERY_SCHEMA = ANALYSIS_SCHEMA


def build_analysis_messages(system_prompt, reference_sql, target_sql,
                             correction_note: str = ""):
    """
    Call 1 of 2: Classification-only Discovery.
    Focuses entirely on: table roles, category extraction, evidence table,
    classification, LEAD partitions. No insertion planning — that competes
    for output tokens and causes the model to cut corners on classification.
    Max output tokens: 4000.
    """
    correction_block = (
        f"\n{correction_note}\n\n" if correction_note else ""
    )
    content_blocks = [
        {
            "type": "text",
            "text": system_prompt,
            "cache_control": {"type": "ephemeral"},
        },
        {
            "type": "text",
            "text": f"REFERENCE SQL:\n{reference_sql}",
            "cache_control": {"type": "ephemeral"},
        },
        {
            "type": "text",
            "text": f"""CRITICAL: RETURN ONLY VALID JSON — NO PROSE, NO MARKDOWN, NO EXPLANATION.
{correction_block}
\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
TASK — ANALYSIS ONLY (no insertion planning in this call)
\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
Read the TARGET SQL and return ONLY the analysis JSON below.
Do NOT include any insertions array. Do NOT determine anchors yet.
Your only job in this call is to read the script and classify it correctly.

STEP 1 — Map functional roles (1A).
  For each temp table identify: FILTER, CHARGE_DETAIL, LINE_PRICING,
  NCCI_BUNDLE, AGGREGATE_DATE, AGGREGATE_ENC, HIERARCHY, TOTAL_PRICE,
  FINAL_JOIN, OUTPUT.

STEP 2 — Extract categories and INDICATOR_FLAGs (1B).
  From LINE_PRICING, extract every service category.
  Separately identify INDICATOR_FLAG columns (binary 0/1 only, never dollar).
  INDICATOR_FLAGs go in indicator_flag_categories ONLY — never in max_categories.

STEP 3 — Build the raw evidence table (1C STEP 1).
  Find AGGREGATE_DATE (GROUP BY EncounterID + ServiceDate).
  Find AGGREGATE_ENC (GROUP BY EncounterID).
  For EVERY column in BOTH steps, record the exact aggregation function
  (MAX or SUM) in aggregate_date_raw.columns and aggregate_enc_raw.columns.
  List every single column. Do not skip any. Do not summarise.

STEP 4 — Classify from the evidence table ONLY (1C STEP 2).
  Apply these rules one column at a time. Use NO other information:
    agg_date=MAX  AND  agg_enc=MAX   \u2192  MAX
    agg_date=SUM  AND  agg_enc=SUM   \u2192  SUM
    agg_date=MAX  AND  agg_enc=SUM   \u2192  HYBRID
    agg_date=SUM  AND  agg_enc=MAX   \u2192  MAX  (unusual — note it)
    binary 0/1 in LINE_PRICING only  \u2192  INDICATOR_FLAG
  Do NOT use category names or prior knowledge. Only the evidence table.

STEP 5 — Self-check (1C STEP 3).
  For each entry in hybrid_categories: confirm agg_date=MAX AND agg_enc=SUM.
    If not \u2192 move to the correct bucket.
  For each entry in sum_categories: confirm agg_date=SUM (or absent).
    If agg_date=MAX AND agg_enc=SUM \u2192 it is HYBRID, move it.
  For each entry in max_categories: confirm agg_enc=MAX.
    If agg_enc=SUM \u2192 wrong, move it.
  Repeat until no entry fails its check.

STEP 5b — Column survival check (1G). MANDATORY.
  Trace how the line-level source table (AGGREGATE_DATE input, e.g. #Step3)
  was built. Determine whether it uses SELECT * or an explicit column list.
  If it uses an EXPLICIT COLUMN LIST, record which columns are present and
  which are absent. This determines whether a JOIN back to the CHARGE_DETAIL
  table is needed in #LineBreakdown for any payment-critical columns.

  Set line_level_has_narrow_column_list = true if the source table uses an
  explicit column list that omits any columns needed for LinePayment.
  If true, identify which payment columns require a JOIN back to the
  CHARGE_DETAIL table (e.g. BillCharges for EPS calculation).

  CRITICAL: A narrow column list NEVER justifies using a pre-LEAD table
  (e.g. #Step2) as the source for #bRanked. The LEAD slot columns ONLY
  exist in the line-level source table. Using an earlier table causes a
  guaranteed runtime crash.

STEP 6 — WINDOW_REDUCTION LEAD partition scope (1E).
  For each LEAD() category: read the PARTITION BY clause exactly.
  Set lead_partition_by, branked_partition_by, bslots_group_by to match.

STEP 7 — Extract contract period boundaries (1H). MANDATORY.
  Scan ALL CASE/WHEN BETWEEN date ranges across every pricing category
  in the LINE_PRICING step. Collect every unique start date and end date.
  Sort chronologically. Define the master period list with no gaps and
  no overlaps. Each period must have a short human-readable label
  (e.g. "Year1", "Year2") and exact date_from / date_to strings
  in YYYY-MM-DD format. These will be used to generate
  ContractEffectiveDateFrom and ContractEffectiveDateTo CASE blocks
  in #LineBreakdown. If a category uses a simple < date threshold
  instead of BETWEEN, infer the boundaries from those thresholds.
  Populate the contract_periods array. It must NOT be empty.
  IMPORTANT: contract_periods is a SEPARATE top-level array — do NOT
  place period objects inside aggregate_date_raw, aggregate_enc_raw,
  max_categories, sum_categories, hybrid_categories, or any other array.

\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
MANDATORY JSON RESPONSE FORMAT
\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
{ANALYSIS_SCHEMA}

TARGET SQL:
{target_sql}

REMINDER: START WITH {{ END WITH }} — NO TEXT OUTSIDE THE JSON.
""",
        },
    ]
    return [{"role": "user", "content": content_blocks}]


def build_planning_messages(system_prompt, reference_sql, target_sql, analysis: dict):
    """
    Call 2 of 2: Insertion planning only.
    Given the validated analysis, determine the four anchor snippets and
    insertion positions. No classification work — just find the anchors.
    Max output tokens: 2000.
    """
    content_blocks = [
        {
            "type": "text",
            "text": system_prompt,
            "cache_control": {"type": "ephemeral"},
        },
        {
            "type": "text",
            "text": f"REFERENCE SQL:\n{reference_sql}",
            "cache_control": {"type": "ephemeral"},
        },
        {
            "type": "text",
            "text": f"""CRITICAL: RETURN ONLY VALID JSON — NO PROSE, NO MARKDOWN, NO EXPLANATION.

\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
TASK — INSERTION PLANNING ONLY
\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
The analysis below has already been validated. Do NOT re-classify anything.
Your only job is to find the four anchor snippets in the TARGET SQL.

VALIDATED ANALYSIS:
{json.dumps(analysis, indent=2)}

For each of the four required insertions, find the anchor in the TARGET SQL:

  1. cleanup_drop
     WHERE: the cleanup / DROP TABLE block at the very TOP of the script
     ANCHOR: exact last ~80 chars of the last existing DROP TABLE line there

  2. breakdown_block
     WHERE: immediately after the HIERARCHY step ({analysis.get('hierarchy_table','#Step4')}),
            before the TOTAL_PRICE step ({analysis.get('total_price_table','#Step5')})
     ANCHOR: exact last ~80 chars of the final SQL statement of the HIERARCHY step

  3. final_join_detail_column
     WHERE: immediately before the INTO clause of the FINAL_JOIN SELECT
     ANCHOR: exact text of the INTO #TableName line (e.g. 'Into #Step7')
     insert_position: "before"

  4. auditor_queries
     WHERE: after the very last SELECT in the script (the OUTPUT step)
     ANCHOR: exact last ~80 chars of the final OUTPUT SELECT statement

ANCHOR RULES:
  - Must be the FINAL statement of the role-identified block, not the opening.
  - Must be unique within the script — not a repeating pattern.
  - Derived from functional role, not assumed table name.

\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
MANDATORY JSON RESPONSE FORMAT
\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
{PLANNING_SCHEMA}

TARGET SQL:
{target_sql}

REMINDER: START WITH {{ END WITH }} — NO TEXT OUTSIDE THE JSON.
""",
        },
    ]
    return [{"role": "user", "content": content_blocks}]


# Legacy wrapper so any external call to build_discovery_messages still works
def build_discovery_messages(system_prompt, reference_sql, target_sql,
                             correction_note: str = ""):
    return build_analysis_messages(system_prompt, reference_sql, target_sql,
                                   correction_note)

# ─────────────────────────────────────────────────────────────────────────────
# Content-generation prompts
# ─────────────────────────────────────────────────────────────────────────────

PHASE_SCHEMAS = {
    "cleanup_drop": """{
  "insertion_id": <int>,
  "phase": "cleanup_drop",
  "anchor_snippet": "<same anchor from discovery>",
  "insert_position": "after",
  "new_content": "IF OBJECT_ID('tempdb..#bRanked') IS NOT NULL DROP TABLE #bRanked\\nIF OBJECT_ID('tempdb..#bSlots') IS NOT NULL DROP TABLE #bSlots\\nIF OBJECT_ID('tempdb..#LineBreakdown') IS NOT NULL DROP TABLE #LineBreakdown\\nIF OBJECT_ID('tempdb..#recon_temp') IS NOT NULL DROP TABLE #recon_temp"
}""",

    "breakdown_block": """{
  "insertion_id": <int>,
  "phase": "breakdown_block",
  "anchor_snippet": "<same anchor from discovery>",
  "insert_position": "after",
  "confirmed_source_columns": ["col1", "col2", "...every column confirmed present in line_level_source_table"],
  "confirmed_output_columns": ["EncounterID", "Sequence", "ProcedureCode", "...every column in the SELECT INTO #LineBreakdown"],
  "new_content": "<ONLY net-new SQL: SELECT INTO #bRanked, optional SELECT INTO #bSlots, SELECT INTO #LineBreakdown>"
}""",

    "final_join_detail_column": """{
  "insertion_id": <int>,
  "phase": "final_join_detail_column",
  "anchor_snippet": "<same anchor from discovery>",
  "insert_position": "after",
  "new_content": "<ONLY the Price_Breakdown_Detail column — starts with , [Price_Breakdown_Detail] = >"
}""",

    "auditor_queries": """{
  "insertion_id": <int>,
  "phase": "auditor_queries",
  "anchor_snippet": "<same anchor from discovery>",
  "insert_position": "after",
  "new_content": "<ONLY the Phase 4 auditor block: comment banner, Query 1, Query 2 reconciliation>"
}""",
}

PHASE_INSTRUCTIONS = {
    "cleanup_drop": """
Generate THREE DROP TABLE safety lines for the new temp tables introduced by
the breakdown block. Output EXACTLY these three lines in this order:

IF OBJECT_ID('tempdb..#bRanked')       IS NOT NULL DROP TABLE #bRanked
IF OBJECT_ID('tempdb..#bSlots')        IS NOT NULL DROP TABLE #bSlots
IF OBJECT_ID('tempdb..#LineBreakdown') IS NOT NULL DROP TABLE #LineBreakdown
IF OBJECT_ID('tempdb..#recon_temp')    IS NOT NULL DROP TABLE #recon_temp
Nothing else in new_content.""",


    "breakdown_block": """
Generate the full line-level breakdown block using TEMP TABLES (not CTEs).

════════════════════════════════════════════════════════════════
COLUMN SURVIVAL VERIFICATION — MANDATORY BEFORE WRITING ANY SQL
════════════════════════════════════════════════════════════════
Before generating any SQL you MUST:

  1. Identify analysis.line_level_source_table (e.g. #Step3Bundle).
  2. Trace its full construction chain in COMPLETE CURRENT SQL STATE above:
       - Find the SELECT statement that creates it.
       - If it uses SELECT *, trace upstream to find that table's column list.
         Keep tracing through consecutive SELECT * steps until you reach an
         explicit column list. Never assume a column is present without
         tracing all the way back to an explicit list.
       - If it uses an explicit column list, that list is the definitive set.
       - Repeat for every intermediate step back to the source table.
  3. Build a confirmed column list — only columns proven present in the
     final source table may be referenced in the SELECT INTO #LineBreakdown.
  4. For any column your breakdown logic needs that is NOT confirmed present,
     use a NULL placeholder: CAST(NULL AS <appropriate_type>) AS [ColumnName]
     NEVER reference a column that is not confirmed present in the source.
  5. Return confirmed_source_columns and confirmed_output_columns in your JSON.
     These are passed to downstream phases so they know exactly which columns
     exist in #LineBreakdown.

This applies to EVERY column — not just Quantity.

════════════════════════════════════════════════════════════════
NARROW COLUMN LIST RULE — MANDATORY
════════════════════════════════════════════════════════════════
If the line-level source table (e.g. #Step3) was built with an EXPLICIT
COLUMN LIST (not SELECT *) that omits columns needed by #LineBreakdown,
you MUST NOT fall back to an earlier table like #Step2.

The line-level source table is ALWAYS the correct source for #bRanked
because it is the ONLY table that contains the LEAD-computed slot columns
(AS1-AS10, Rad1-Rad10, etc.). These columns do not exist in #Step2 or
any earlier table. Using #Step2 as the source for #bRanked causes a
guaranteed runtime crash on #bSlots: "Invalid column name 'AS1'".

Handle missing columns from the narrow source table as follows:

  (a) INFORMATIONAL COLUMNS (BillCharges, Plan, Payer01Code, AGE, etc.
      — columns used for display only, not for payment math):
      Use a NULL placeholder in the SELECT INTO #LineBreakdown:
        CAST(NULL AS VARCHAR(50)) AS [Plan]
        CAST(NULL AS DECIMAL(12,2)) AS [BillCharges]
      These are safe to NULL because they do not affect LinePayment.

  (b) PAYMENT COLUMNS genuinely required for LinePayment calculation
      (e.g. BillCharges needed for EPS: Round(BillCharges * 0.74, 2)):
      Add a LEFT JOIN back to the CHARGE_DETAIL table (#Step2 or equivalent)
      inside the SELECT INTO #LineBreakdown using EncounterID + Sequence
      as the join key. Only JOIN for columns that are truly required for
      the payment formula — do not JOIN just to recover display columns.

      Example:
        FROM #bRanked b
        INNER JOIN #Step4 s4 ON s4.[EncounterID] = b.[EncounterID]
        LEFT  JOIN #bSlots rd ON rd.[EncounterID] = b.[EncounterID]
        LEFT  JOIN #Step2  src ON src.[EncounterID] = b.[EncounterID]
                               AND src.[Sequence]   = b.[Sequence]

DO NOT recompute slot values (AS1-AS10, Rad1-Rad10) from scratch using
division logic. The slot values are already correctly computed by the
LEAD() expressions in the source table. Read them directly.

════════════════════════════════════════════════════════════════
GENERATION RULES — USE TEMP TABLES, NOT CTEs
════════════════════════════════════════════════════════════════
Generate THREE sequential temp table blocks:

BLOCK 1 — #bRanked (replaces bRanked CTE):
  SELECT *
      -- MAX category example — one ROW_NUMBER per MAX category, named rn_ + FULL category name:
      , ROW_NUMBER() OVER (PARTITION BY [EncounterID]
            ORDER BY [ER] DESC, [Amount] DESC, [Sequence] ASC) AS rn_ER
      , ROW_NUMBER() OVER (PARTITION BY [EncounterID]
            ORDER BY [Gamma_Knife] DESC, [Amount] DESC, [Sequence] ASC) AS rn_Gamma_Knife
      -- HYBRID category example — PARTITION BY EncounterID+ServiceDate:
      , ROW_NUMBER() OVER (PARTITION BY [EncounterID], [ServiceDate]
            ORDER BY [Cardiac_Cath] DESC, [Amount] DESC, [Sequence] ASC) AS rn_Cardiac_Cath
      -- WINDOW_REDUCTION example — PARTITION matches the LEAD partition scope:
      --   Check analysis.window_reduction_lead_partitions[category].branked_partition_by.
      --   ORDER BY the SUM of ALL slot columns for that category (never a single slot).
      , ROW_NUMBER() OVER (PARTITION BY [EncounterID]
            ORDER BY (ISNULL([AS1],0)+ISNULL([AS2],0)+ISNULL([AS3],0)+ISNULL([AS4],0)+
                      ISNULL([AS5],0)+ISNULL([AS6],0)+ISNULL([AS7],0)+ISNULL([AS8],0)+
                      ISNULL([AS9],0)+ISNULL([AS10],0)) DESC,
                     [Amount] DESC, [Sequence] ASC) AS rn_AS
      -- Repeat for every MAX, HYBRID, WINDOW_REDUCTION, and INDICATOR_FLAG category.
      -- COMPLETENESS CHECK: count of ROW_NUMBER() lines must equal
      --   (# MAX) + (# HYBRID) + (# WINDOW_REDUCTION) + (# INDICATOR_FLAG).
  INTO #bRanked
  FROM [line_level_source_table]   -- e.g. #Step3Bundle — NEVER #Step2 or earlier

BLOCK 2 — #bSlots (replaces bSlots CTE — ONLY if analysis.has_window_reduction is true):
  SELECT
      [EncounterID]
      -- Include [ServiceDate] in GROUP BY ONLY if the LEAD for that category
      -- is partitioned by EncounterID+ServiceDate.
      -- Check analysis.window_reduction_lead_partitions[category].bslots_group_by:
      --   "EncounterID_only"            => GROUP BY [EncounterID]
      --   "EncounterID_and_ServiceDate" => GROUP BY [EncounterID], [ServiceDate]
      -- The #bSlots GROUP BY MUST match the LEAD partition scope.
      -- Do NOT default to GROUP BY [EncounterID],[ServiceDate] for all scripts.
      -- Example using real names — replace rn_AS with the actual WINDOW_REDUCTION alias:
      , MAX(CASE WHEN rn_AS = 1 THEN ISNULL([AS1], 0) ELSE 0 END) AS AS1
      , MAX(CASE WHEN rn_AS = 1 THEN ISNULL([AS2], 0) ELSE 0 END) AS AS2
      -- ... repeat for all slot columns using the actual rn_ alias defined in #bRanked
  INTO #bSlots
  FROM #bRanked
  GROUP BY [EncounterID]  -- add [ServiceDate] only if bslots_group_by = "EncounterID_and_ServiceDate"
  -- Use GROUP BY + MAX(CASE...) instead of WHERE rn=1 OR rn=1
  -- This correctly handles multiple window categories in one pass.
  If analysis.has_window_reduction is false, omit #bSlots entirely.

BLOCK 3 — #LineBreakdown (the main output):
  SELECT DISTINCT
      b.[EncounterID], b.[Sequence], b.[ProcedureCode], b.[RevenueCode],
      b.[ServiceDate], b.[Amount] AS [BilledAmount],
      Quantity (NULL placeholder CAST(NULL AS NUMERIC(18,4)) if not confirmed present),
      [ServiceCategory] = CASE ... END   (full hierarchy CASE),
      [PricingMethod]   = CASE ... END   (human-readable rate description),
      [LinePayment]     = ROUND(...)     (MAX/SUM/HYBRID/slot-pivot logic),
      [BundledByNCCI]   = CAST(b.[NeedsCorrectedClaim] AS TINYINT)
                          OR CAST(0 AS TINYINT) if no NCCI bundle
      [ContractEffectiveDateFrom] = CASE
          -- One WHEN branch per period from analysis.contract_periods:
          WHEN b.[ServiceDate] BETWEEN 'date_from' AND 'date_to' THEN CAST('date_from' AS DATE)
          -- ... repeat for every period, then ELSE NULL
          ELSE NULL END
      [ContractEffectiveDateTo] = CASE
          -- Mirror of ContractEffectiveDateFrom — same BETWEEN ranges:
          WHEN b.[ServiceDate] BETWEEN 'date_from' AND 'date_to' THEN CAST('date_to' AS DATE)
          -- ... repeat for every period, then ELSE NULL
          ELSE NULL END
      -- MANDATORY: use EVERY period from analysis.contract_periods. No gaps.
  INTO #LineBreakdown
  FROM #bRanked b
  INNER JOIN [hierarchy_table] s4 ON s4.[EncounterID] = b.[EncounterID]
  LEFT  JOIN #bSlots rd ON rd.[EncounterID] = b.[EncounterID]
             -- CRITICAL: if bslots_group_by = "EncounterID_and_ServiceDate",
             -- also add: AND rd.[ServiceDate] = b.[ServiceDate]
             -- If bslots_group_by = "EncounterID_only", join on EncounterID alone.
             -- Match the JOIN key to whatever #bSlots is grouped by.
  -- Omit #bSlots join entirely if has_window_reduction is false
  ORDER BY b.[EncounterID], b.[Sequence]

ServiceCategory labeling rules by category type:
  MAX:     IIF(b.rn_ER = 1, 'ER', 'ER_Non_Winner')   -- use the actual rn_ alias for each category
           One winner per encounter. rn>1 = Non_Winner.
  SUM:     'CategoryName' (flat label, all matching lines pay, no rank check)
  HYBRID:  IIF(b.rn_Cardiac_Cath = 1, 'Cardiac_Cath', 'Cardiac_Cath_Non_Winner')   -- actual alias
           HYBRID uses the SAME labeling as MAX — because the PARTITION is
           EncounterID+ServiceDate, rn=1 means winner FOR THAT DATE only.
           Lines with rn>1 on any given date are Non_Winners for that date.
           Multiple rn=1 rows (one per date) all receive the category name.
           DO NOT label all HYBRID lines with the flat category name (SUM
           behavior) — that would obscure which lines actually drove payment.

IMPORTANT NOTES:
  • No CTEs — use temp tables throughout.
  • No semicolon rule needed — temp tables use plain SELECT INTO, no WITH keyword.
  • Do NOT rewrite or repeat any existing SQL.
  • ServiceCategory CASE must cover every category in exact hierarchy order.
  • OP_Default is always last dollar category. INDICATOR_FLAG categories come after.
  • rn_ NAMING CONVENTION — NON-NEGOTIABLE: Every ROW_NUMBER() alias MUST be
    rn_ + FULL category name, NO abbreviations. rn_Chemotherapy not rn_Chemo,
    rn_CardiacCath not rn_CC, rn_GammaKnife not rn_GK. This alias is used
    identically in #bRanked (definition) AND #LineBreakdown (reference). Any
    mismatch or abbreviation causes "Invalid column name" SQL errors at runtime.
  • INDICATOR_FLAG categories (binary 0/1, no dollar value — e.g. COVID, CPT flags):
      - Must have ROW_NUMBER() in #bRanked: PARTITION BY [EncounterID], same as MAX
      - ServiceCategory: IIF(b.rn_FLAG = 1, 'FLAG_Name', 'FLAG_Name_Non_Winner')
        placed AFTER all dollar categories (OP_Default and Suppressed_By_Hierarchy)
      - LinePayment: $0 contribution — DO NOT add flag value to LinePayment sum.
        A binary 1 added to LinePayment creates $1 phantom payments breaking recon.
      - Suppressed_By_Hierarchy sum: MUST include ISNULL(b.[FLAG],0) so pure-flag
        lines get 'Suppressed_By_Hierarchy' not 'No_Payment_Category'.
  • LinePayment rules by category type:
      MAX:              IIF(s4.[ER] != 0, IIF(b.rn_ER = 1, ISNULL(b.[ER], 0), 0), 0)  -- use actual names
      SUM:              IIF(s4.[CAT] != 0, ISNULL(b.[CAT], 0), 0)
      HYBRID:           IIF(s4.[Cardiac_Cath] != 0, IIF(b.rn_Cardiac_Cath = 1, ISNULL(b.[Cardiac_Cath], 0), 0), 0)
                        HYBRID uses IIF(rn=1) — identical pattern to MAX. Because rn is
                        partitioned by EncounterID+ServiceDate, every date's winner (rn=1)
                        contributes its value; rn>1 rows on the same date contribute $0.
                        DO NOT use plain SUM logic for HYBRID — that pays every line and
                        produces an incorrect encounter total.
      INDICATOR_FLAG:   $0 — no IIF, no contribution to LinePayment sum whatsoever.
                        The flag value is binary 0 or 1; adding it creates $1 phantom payments.
      WINDOW_REDUCTION: slot-pivot from #bSlots, keyed by EncounterID alone OR
                        EncounterID+ServiceDate — match the #bSlots JOIN key used above.

════════════════════════════════════════════════════════════════
CATEGORIES WITH SUB-COLUMNS — CRITICAL SPECIAL HANDLING
════════════════════════════════════════════════════════════════
Some categories are composed of multiple sub-columns at the line level that
get merged into a single column in the encounter-level hierarchy table (#Step4).
You MUST check for this pattern and handle it correctly.

HOW TO DETECT:
  In Phase 1B (LINE_PRICING step), look for any category that has more than
  one column contributing to its total. Common pattern: a main category column
  plus one or more named sub-columns (e.g. [Implants], [Implant_V2790],
  [implant_C1713] are three separate columns in LINE_PRICING that #Step4 merges
  into a single [Implants] column). Check which of these sub-columns survive
  into the line-level source table via the column survival trace.

RULE FOR ServiceCategory:
  When a category has sub-columns, the ServiceCategory WHEN condition must check
  the SUM of all confirmed sub-columns, not just the main column alone:
    WHEN s4.[Implants] != 0
     AND (ISNULL(b.[Implants],0) + ISNULL(b.[Implant_V2790],0) + ISNULL(b.[implant_C1713],0)) > 0
        THEN 'Implants'
  A line with b.[Implants]=0 but b.[Implant_V2790]>0 would incorrectly fall
  through to Suppressed_By_Hierarchy if only b.[Implants] is checked.

RULE FOR LinePayment:
  Similarly, LinePayment must sum all confirmed sub-columns:
    + IIF(s4.[Implants] != 0,
          ISNULL(b.[Implants],0) + ISNULL(b.[Implant_V2790],0) + ISNULL(b.[implant_C1713],0),
          0)
  Using only b.[Implants] will silently produce $0 for V2790/C1713 lines,
  causing reconciliation failures.

DO NOT RE-IMPLEMENT #Step4 SUPPRESSION LOGIC:
  s4.[Implants] already reflects the correct encounter-level suppression outcome
  from #Step4 (including any complex CASE/IIF conditions). If s4.[Implants] != 0,
  the category survived for that encounter — regardless of what other categories
  also survived. Do NOT re-derive suppression by checking s4.[Minor_Surgery],
  s4.[ER], etc. in the ServiceCategory CASE for Implants. Trust s4 as the
  authoritative post-suppression source.""",

    "final_join_detail_column": """Generate ONLY the Price_Breakdown_Detail correlated subquery column per
Phase 3 of the system prompt.

CONFIRMED #LineBreakdown COLUMNS (carry-forward from breakdown_block phase):
{confirmed_lb_columns_block}

CRITICAL: Only reference columns from the confirmed list above when querying
#LineBreakdown. Any column not in that list does not exist in #LineBreakdown
and will cause a SQL error.

  • This is the LAST column in the FINAL_JOIN SELECT.
  • The anchor is the INTO #TableName line itself (e.g. "Into #Step7").
  • insert_position is "before" — the column lands immediately before INTO.
  • Do NOT anchor to the last column definition — anchor to the INTO line.
  • Use STRING_AGG(… , ' || ') WITHIN GROUP (ORDER BY lb.Sequence).
  • Reference #LineBreakdown as lb.
  • Match the encounter alias used in the existing FINAL_JOIN step
    (inspect CURRENT SQL STATE above to find it).
  • Output ONLY the column expression — must start with:
      , [Price_Breakdown_Detail] =
  • Do NOT include surrounding SELECT, INTO, or FROM clauses.""",


    "auditor_queries": """Generate the full Phase 4 auditor block per the system prompt.

CONFIRMED #LineBreakdown COLUMNS (carry-forward from breakdown_block phase):
{confirmed_lb_columns_block}

CRITICAL: All queries may ONLY reference columns from the confirmed list above
when selecting from #LineBreakdown. Any column not in that list does not exist
and will cause a SQL error.

Generate the following four items in order:

  1. Comment banner (write EXACTLY this, replacing the placeholder with the word QUERIES):
     --======================================================================
     -- *** [AUDITOR PRICE BREAKDOWN SUMMARY <PLACEHOLDER> - INSERTED SECTION] ***
     --======================================================================
     (replace <PLACEHOLDER> with: QUERIES)

  2. Query 1 — Full line-level detail SELECT:
     Standard confirmed columns to include (omit any not in confirmed list):
       lb.[EncounterID], lb.[Sequence], lb.[ProcedureCode], lb.[RevenueCode],
       lb.[ServiceDate], lb.[BilledAmount], lb.[Quantity],
       lb.[ContractEffectiveDateFrom], lb.[ContractEffectiveDateTo],
       lb.[ServiceCategory], lb.[PricingMethod],
       lb.[LinePayment] AS [ContractualPayment_Line],
       lb.[BundledByNCCI] AS [ZeroedByNCCI_Flag]

  3. INSERT INTO persistent table — immediately after Query 1 SELECT:
     Insert the same line-level results into [Automation].[dbo].[LineBreakdown_Results].
     Include a ScriptName column derived from the input filename (use the actual
     filename being processed, which can be inferred from the script header comments
     or from the USE [database] statement at the top — use whatever identifier best
     describes this specific contract script).

     INSERT INTO [Automation].[dbo].[LineBreakdown_Results]
     (EncounterID, Sequence, ProcedureCode, RevenueCode, ServiceDate,
      BilledAmount, Quantity, ContractEffectiveDateFrom, ContractEffectiveDateTo,
      ServiceCategory, PricingMethod,
      LinePayment, BundledByNCCI, ScriptName, RunDate)
     SELECT
         lb.[EncounterID], lb.[Sequence], lb.[ProcedureCode], lb.[RevenueCode],
         lb.[ServiceDate], lb.[BilledAmount], lb.[Quantity],
         lb.[ContractEffectiveDateFrom], lb.[ContractEffectiveDateTo],
         lb.[ServiceCategory], lb.[PricingMethod],
         lb.[LinePayment], lb.[BundledByNCCI],
         '<ScriptIdentifier>',
         GETDATE()
     FROM #LineBreakdown lb

     Only include columns in the INSERT that are confirmed present in #LineBreakdown.
     If Quantity is a NULL placeholder it is still inserted (NULL is valid).

  4. Query 2 — Reconciliation:
     DROP TABLE IF EXISTS #recon_temp
     SELECT EncounterID, SUM(LinePayment) AS LineBreakdown_Total
     INTO #recon_temp FROM #LineBreakdown GROUP BY EncounterID
     Then discrepancy SELECT joining #recon_temp to analysis.total_price_table
     WHERE ABS(LineBreakdown_Total - Price) > 1.00
     ORDER BY ABS discrepancy DESC.

Rules:
  • Use the ACTUAL table names from the analysis for all references.
  • Do NOT repeat or rewrite any existing SQL — purely additive.""",
}




def build_content_messages(system_prompt, reference_sql, current_sql,
                           context, analysis, insertion,
                           confirmed_lb_columns=None):
    """
    confirmed_lb_columns: list of column names confirmed present in #LineBreakdown.
    Populated after breakdown_block phase and passed to subsequent phases so they
    never reference a column that doesn't exist in #LineBreakdown.
    """
    phase        = insertion["phase"]
    anchor       = insertion["anchor_snippet"]
    insertion_id = insertion["insertion_id"]

    schema       = PHASE_SCHEMAS.get(phase, """{
  "insertion_id": <int>,
  "phase": "<phase>",
  "anchor_snippet": "<same anchor>",
  "insert_position": "after",
  "new_content": "<net-new SQL only>"
}""")
    instructions = PHASE_INSTRUCTIONS.get(phase, (
        "Generate ONLY the net-new SQL content needed for this insertion. "
        "Do NOT rewrite or repeat any existing SQL."
    ))

    # Inject confirmed #LineBreakdown columns into downstream phase instructions
    if confirmed_lb_columns and "{confirmed_lb_columns_block}" in instructions:
        col_block = "\n".join(f"  - {c}" for c in confirmed_lb_columns)
        instructions = instructions.replace("{confirmed_lb_columns_block}", col_block)
    elif "{confirmed_lb_columns_block}" in instructions:
        instructions = instructions.replace(
            "{confirmed_lb_columns_block}",
            "  (not yet available — breakdown_block phase has not completed)"
        )

    # Extract retry note if present (set when rn_ consistency check fails)
    retry_note_text = insertion.get("_retry_note", "")
    retry_note = (
        f"\n⚠️ RETRY INSTRUCTION — READ THIS FIRST:\n{retry_note_text}\n\n"
        if retry_note_text else ""
    )

    content_blocks = [
        {
            "type": "text",
            "text": system_prompt,
            "cache_control": {"type": "ephemeral"},
        },
        {
            "type": "text",
            "text": f"REFERENCE SQL:\n{reference_sql}",
            "cache_control": {"type": "ephemeral"},
        },
        {
            "type": "text",
            "text": f"""CRITICAL: RETURN ONLY VALID JSON — NO PROSE, NO MARKDOWN, NO EXPLANATION.

════════════════════════════════════════════════════════════════
SCRIPT ANALYSIS  (from discovery — treat as authoritative)
════════════════════════════════════════════════════════════════
{json.dumps(analysis, indent=2)}

════════════════════════════════════════════════════════════════
COMPLETE CURRENT SQL STATE
(verify column names, table aliases, existing structure here)
════════════════════════════════════════════════════════════════
{current_sql}

════════════════════════════════════════════════════════════════
INSERTION TASK  —  insertion_id={insertion_id}  phase={phase}
════════════════════════════════════════════════════════════════
{retry_note}{instructions}

LOCAL CONTEXT AROUND INSERTION POINT:
{context}

ANCHOR (insertion happens immediately after this):
{anchor}

════════════════════════════════════════════════════════════════
MANDATORY JSON RESPONSE FORMAT
════════════════════════════════════════════════════════════════
{schema}

REMINDER: new_content = NET-NEW SQL ONLY.  START WITH {{ END WITH }}.
""",
        },
    ]
    return [{"role": "user", "content": content_blocks}]


# ─────────────────────────────────────────────────────────────────────────────
# Claude API call with retry
# ─────────────────────────────────────────────────────────────────────────────

def call_claude_with_retry(client, messages, label, max_retries=3, max_tokens=40000):
    global total_cached_tokens, total_uncached_tokens

    for attempt in range(max_retries + 1):
        try:
            print(f"\n  ── {label} (attempt {attempt + 1}) ──")
            t0 = time.perf_counter()

            text = ""
            msg = None
            
            try:
                msg  = client.messages.create(
                    model=MODEL, max_tokens=max_tokens, messages=messages
                )
                text = "".join(
                    b.text for b in msg.content
                    if getattr(b, "type", None) == "text"
                )
                print(f"  Standard response received ({len(text)} chars)")
                
            except Exception as e:
                if "Streaming is required" not in str(e):
                    raise
                    
                print("  Switching to streaming…")
                text = ""
                stream_chunks = 0
                
                try:
                    with client.messages.stream(
                        model=MODEL, max_tokens=max_tokens, messages=messages
                    ) as stream:
                        for chunk in stream:
                            if (chunk.type == "content_block_delta"
                                    and hasattr(chunk.delta, "text")):
                                text += chunk.delta.text
                                stream_chunks += 1
                                
                                # Progress indicator for large streams
                                if stream_chunks % 100 == 0:
                                    print(f"    Streaming... {len(text)} chars received")
                        
                        msg = stream.get_final_message()
                        print(f"  Streaming complete ({stream_chunks} chunks, {len(text)} chars)")
                        
                except Exception as stream_error:
                    print(f"  Streaming error: {stream_error}")
                    raise

            # Only attempt JSON parsing after we have the complete response
            if not text.strip():
                raise ValueError("Empty response received from Claude")

            elapsed = time.perf_counter() - t0
            
            # Parse JSON from complete response
            parsed = extract_json(text)

            # Check if parsing failed and we should retry
            if "error" in parsed and attempt < max_retries:
                error_msg = parsed.get("error", "Unknown JSON error")
                print(f"  JSON parse failed: {error_msg}")
                print(f"  Response length: {len(text)} chars")
                print(f"  Response preview: {text[:200]}...")
                print(f"  Retrying in 1.5s…")
                time.sleep(1.5)
                continue

            # Calculate usage and cost
            inp, out, _, cache_read = usage_numbers(msg) if msg else (0, 0, 0, 0)
            total_cached_tokens   += cache_read
            total_uncached_tokens += inp
            cost = estimate_cost(inp, out, cache_read)

            print(f"  in={inp:,} out={out:,} cached={cache_read:,} "
                  f"cost=${cost:.4f} time={elapsed:.1f}s")
            
            return parsed, msg, elapsed, cost

        except Exception as e:
            err = str(e)
            print(f"  API error: {err}")
            if attempt < max_retries:
                wait = (2 ** attempt + random.uniform(0, 1)) if "429" in err else 1.5
                print(f"  Retrying in {wait:.1f}s…")
                time.sleep(wait)
            else:
                raise


# ─────────────────────────────────────────────────────────────────────────────
# Required-phase validation
# ─────────────────────────────────────────────────────────────────────────────

def check_required_phases(insertions: list, log_path) -> list[str]:
    found   = {ins.get("phase") for ins in insertions}
    missing = sorted(REQUIRED_PHASES - found)
    if missing:
        msg = f"Missing required phases: {missing}"
        print(f"\n  WARNING: {msg}")
        append_log(log_path, f"MISSING_PHASES | {msg}")
    return missing


# ─────────────────────────────────────────────────────────────────────────────
# Single-file processor
# ─────────────────────────────────────────────────────────────────────────────

def process_single_file(client, input_file_path, system_prompt,
                        reference_sql, file_outputs, session_id):
    global db_manager

    print(f"\n{'='*70}")
    print(f"  Processing: {input_file_path.name}")
    print(f"{'='*70}")

    file_id    = None
    file_start = time.perf_counter()

    if db_manager:
        try:
            file_id = db_manager.start_file(
                session_id=session_id,
                filename=input_file_path.name,
                input_file_path=str(input_file_path),
                file_size_bytes=input_file_path.stat().st_size,
            )
        except Exception as e:
            print(f"  DB error starting file: {e}")

    insertions_log = {
        "session_info": {
            "timestamp":  datetime.now().isoformat(),
            "input_file": str(input_file_path),
            "approach":   "Role-based anchor insertion v9 — INDICATOR_FLAG + rn_ naming + consistency retry",
            "file_id":    file_id,
            "session_id": session_id,
        },
        "analysis":       {},
        "missing_phases": [],
        "insertions":     [],
    }

    file_total_input  = 0
    file_total_output = 0
    file_total_time   = 0.0
    file_total_cost   = 0.0

    try:
        target_sql = read_text(input_file_path)

        # ── Phase 1: Discovery + analysis ────────────────────────────────────
        # ── Phase 1a: Analysis (classification only, max 6000 tokens) ─────────
        print("\n  Phase 1a — Analysis (classification)…")
        analysis_msgs = build_analysis_messages(system_prompt, reference_sql, target_sql)
        analysis_resp, msg, elapsed, cost = call_claude_with_retry(
            client, analysis_msgs, "ANALYSIS",
            max_tokens=6000,
        )

        inp, out, _, _ = usage_numbers(msg)
        file_total_input  += inp
        file_total_output += out
        file_total_time   += elapsed
        file_total_cost   += cost

        if "error" in analysis_resp:
            raise RuntimeError(f"Analysis JSON parse failed: {analysis_resp['error']}")

        analysis = analysis_resp  # analysis call returns the analysis dict directly

        # insertions not available until after Phase 1b (planning call).
        # missing_phases check is deferred until after the retry loop.
        insertions    = []   # populated by Phase 1b below
        missing_phases = []  # populated after Phase 1b

        # ── Classification validation with fatal-error retry loop ────────────
        # Run validate_hybrid_classification against the Discovery output.
        # Fatal issues (HYBRID_MISS, HYBRID_WRONG, SUM_IS_HYBRID, SUM_IS_MAX,
        # MAX_IS_SUM, INDICATOR_IN_MAX) block breakdown_block generation and
        # trigger a Discovery retry with the specific errors injected so the
        # model can correct only the misclassified categories.
        # Non-fatal issues (MAX_IS_HYBRID, LEAD/BSLOTS MISMATCH) are logged
        # and printed but do not block generation.
        MAX_CLASSIFICATION_RETRIES = 3
        classification_attempt     = 0
        all_issues                 = []

        while True:
            all_issues, agg_date_raw, agg_enc_raw = validate_hybrid_classification(analysis)
            fatal_issues   = [i for i in all_issues if i["fatal"]]
            warning_issues = [i for i in all_issues if not i["fatal"]]

            # Print all issues regardless
            if all_issues:
                print(f"\n  Classification check: "
                      f"{len(fatal_issues)} fatal, {len(warning_issues)} warning(s)")
                for iss in all_issues:
                    tag = "✗ FATAL" if iss["fatal"] else "⚠ WARN "
                    print(f"    [{tag}] {iss['code']} | {iss['message']}")
                    append_log(file_outputs["log_file"],
                               f"{'FATAL' if iss['fatal'] else 'WARN'} | "
                               f"{iss['code']} | {iss['message']}")
            else:
                print("  ✓ Classification validated — no issues")

            # No fatal issues → proceed
            if not fatal_issues:
                break

            # Fatal issues but retries exhausted → abort this file
            if classification_attempt >= MAX_CLASSIFICATION_RETRIES:
                msg = (f"Classification still has {len(fatal_issues)} fatal issue(s) "
                       f"after {MAX_CLASSIFICATION_RETRIES} retries — aborting file.")
                print(f"  ✗ {msg}")
                append_log(file_outputs["log_file"], f"CLASSIFICATION_ABORT | {msg}")
                raise RuntimeError(msg)

            classification_attempt += 1

            if classification_attempt == 1:
                # ── Attempt 1: deterministic Python patch — no API call ────────
                # The model already returned correct raw evidence but placed
                # categories in wrong buckets. Rewrite the buckets directly
                # from agg_date_raw / agg_enc_raw. This avoids the oscillation
                # pattern where retrying the LLM fixes flagged errors but
                # breaks previously-correct categories on each pass.
                print(f"\n  Attempt 1: patching classification directly from "
                      f"raw evidence (no API call)…")
                analysis = patch_analysis_classifications(
                    analysis, agg_date_raw, agg_enc_raw
                )
                print(f"  ✓ Patched: max={analysis.get('max_categories',[])} "
                      f"hybrid={analysis.get('hybrid_categories',[])} "
                      f"sum(count)={len(analysis.get('sum_categories',[]))}")

            else:
                # ── Attempt 2+: LLM retry with fully-prescribed corrections ───
                # If the deterministic patch still fails (e.g. evidence was
                # wrong in the first call), fall back to an LLM retry that
                # explicitly states the required bucket for every category.
                print(f"\n  Attempt {classification_attempt}: LLM retry to fix "
                      f"{len(fatal_issues)} remaining fatal issue(s)…")

                # Build exact per-category correction instructions
                corrections_by_cat: dict = {}
                for iss in fatal_issues:
                    m = re.search(r"'([^']+)'", iss["message"])
                    cat = m.group(1) if m else "unknown"
                    corrections_by_cat.setdefault(cat, []).append(iss)

                def _correct_bucket(cat_lower: str) -> str:
                    d = agg_date_raw.get(cat_lower, "")
                    e = agg_enc_raw.get(cat_lower, "")
                    if d == "MAX" and e == "SUM": return "hybrid_categories"
                    if e == "MAX":               return "max_categories"
                    return "sum_categories"

                correction_lines = []
                for cat, issues_for_cat in corrections_by_cat.items():
                    d_fn = agg_date_raw.get(cat, "NOT FOUND IN EVIDENCE")
                    e_fn = agg_enc_raw.get(cat,  "NOT FOUND IN EVIDENCE")
                    correct = _correct_bucket(cat)
                    correction_lines.append(
                        f"  CATEGORY '{cat}':\n"
                        f"    Evidence: AGGREGATE_DATE={d_fn}, AGGREGATE_ENC={e_fn}\n"
                        f"    REQUIRED: move to {correct}. Remove from all other buckets."
                    )

                correction_text = '\n'.join(correction_lines)
                correction_note = (
                    f"CLASSIFICATION CORRECTION — attempt {classification_attempt}\n"
                    f"The previous response still had fatal errors.\n"
                    f"Apply ONLY these changes and leave everything else untouched:\n"
                    f"\n"
                    f"{correction_text}\n"
                    f"\n"
                    f"Return the complete corrected JSON."
                )

                retry_msgs = build_analysis_messages(
                    system_prompt, reference_sql, target_sql,
                    correction_note=correction_note,
                )
                analysis_retry, msg_obj, elapsed2, cost2 = call_claude_with_retry(
                    client, retry_msgs, f"ANALYSIS RETRY {classification_attempt}",
                    max_tokens=6000,
                )
                inp2, out2, _, _ = usage_numbers(msg_obj)
                file_total_input  += inp2
                file_total_output += out2
                file_total_time   += elapsed2
                file_total_cost   += cost2

                if "error" in analysis_retry:
                    raise RuntimeError(
                        f"Analysis retry {classification_attempt} parse failed: "
                        f"{analysis_retry['error']}"
                    )
                analysis = analysis_retry

        # ── Phase 1b: Insertion planning (anchors only, max 2000 tokens) ─────
        # Classification is now validated. Run a separate focused call to find
        # the four anchor snippets. This call gets the model's full attention
        # without competing with classification work.
        print("\n  Phase 1b — Insertion planning (anchor identification)…")
        planning_msgs = build_planning_messages(
            system_prompt, reference_sql, target_sql, analysis
        )
        planning_resp, msg_p, elapsed_p, cost_p = call_claude_with_retry(
            client, planning_msgs, "PLANNING",
            max_tokens=2000,
        )
        inp_p, out_p, _, _ = usage_numbers(msg_p)
        file_total_input  += inp_p
        file_total_output += out_p
        file_total_time   += elapsed_p
        file_total_cost   += cost_p

        if "error" in planning_resp:
            raise RuntimeError(f"Planning JSON parse failed: {planning_resp['error']}")

        insertions = planning_resp.get("insertions", [])

        # Now check required phases
        missing_phases = check_required_phases(insertions, file_outputs["log_file"])
        insertions_log["missing_phases"] = missing_phases

        # Save analysis + plan for audit (after all retries)
        insertions_log["analysis"] = analysis
        insertions_log["classification_issues"] = all_issues
        with open(file_outputs["analysis_log"], "w", encoding="utf-8") as f:
            json.dump({
                "analysis":               analysis,
                "insertion_plan":         insertions,
                "missing_phases":         missing_phases,
                "classification_issues":  all_issues,
                "classification_retries": classification_attempt,
            }, f, indent=2)

        print(f"\n  Analysis:")
        print(f"    line_level_source  : {analysis.get('line_level_source_table', '?')}")
        print(f"    hierarchy_table    : {analysis.get('hierarchy_table', '?')}")
        print(f"    total_price_table  : {analysis.get('total_price_table', '?')}")
        print(f"    final_join_table   : {analysis.get('final_join_table', '?')}")
        print(f"    output_table       : {analysis.get('output_table', '?')}")
        print(f"    has_ncci_bundle    : {analysis.get('has_ncci_bundle', '?')}")
        print(f"    has_window_reduct  : {analysis.get('has_window_reduction', '?')}")
        print(f"    has_agg_date_step  : {analysis.get('has_aggregate_date_step', '?')}")
        print(f"    max_categories     : {analysis.get('max_categories', [])}")
        print(f"    hybrid_categories  : {analysis.get('hybrid_categories', [])}")
        print(f"    indicator_flags    : {analysis.get('indicator_flag_categories', [])}")
        if analysis.get("window_reduction_lead_partitions"):
            print(f"    lead_partitions    :")
            for lp in analysis["window_reduction_lead_partitions"]:
                print(f"      {lp.get('category','?')}: lead={lp.get('lead_partition_by','?')} "
                      f"branked={lp.get('branked_partition_by','?')} "
                      f"bslots={lp.get('bslots_group_by','?')}")
        print(f"    semicolon_needed   : {analysis.get('semicolon_needed_before_cte', '?')}")
        print(f"    insertions planned : {len(insertions)}")
        if missing_phases:
            print(f"    ⚠  MISSING phases  : {missing_phases}")
        fatal_count = sum(1 for i in all_issues if i["fatal"])
        if fatal_count:
            print(f"    ⚠  Classification retries used: {classification_attempt}")

        if not insertions:
            print("  No insertions planned — saving original.")
            file_outputs["patched_sql"].write_text(target_sql, encoding="utf-8")
            _save_log(file_outputs, insertions_log, input_file_path,
                      0, 0, 0, file_total_input, file_total_output,
                      file_total_time, file_total_cost)
            _end_db_file(file_id, "SUCCESS", file_outputs, 0, 0,
                         file_total_cost, time.perf_counter() - file_start)
            return _result(input_file_path.name, True, 0, 0,
                           file_total_cost, time.perf_counter() - file_start)

        # ── Phase 2: Generate + validate + apply ─────────────────────────────
        print(f"\n  Phase 2 — Generating {len(insertions)} insertion(s)…")

        current_sql = target_sql
        successful  = 0
        failed      = 0

        confirmed_lb_columns = []  # populated after breakdown_block, passed to downstream phases
        for i, ins in enumerate(insertions):
            attempt_count = 0  # tracks retries for breakdown_block rn_ consistency
            insertion_id = ins.get("insertion_id", i + 1)
            phase        = ins.get("phase", "unknown")
            description  = ins.get("description", "")
            anchor       = ins.get("anchor_snippet", "")
            position     = ins.get("insert_position", "after")
            role_table   = ins.get("role_anchor_table", "")

            print(f"\n  [{i+1}/{len(insertions)}] phase={phase}")
            print(f"  {description}")

            # ── Anchor validation ─────────────────────────────────────────────
            if not anchor:
                _skip(file_id, insertion_id, description, anchor, position,
                      "Empty anchor_snippet", file_outputs)
                failed += 1
                continue

            anchor_pos = find_anchor(current_sql, anchor)
            if anchor_pos is None:
                _skip(file_id, insertion_id, description, anchor, position,
                      "Anchor not found in current SQL", file_outputs,
                      preview=anchor[:100])
                failed += 1
                continue

            # Compute anchor line number BEFORE insertion
            anchor_line = get_line_number(current_sql, anchor_pos)
            print(f"  Anchor found at line {anchor_line}")

            # ── Idempotency guard ─────────────────────────────────────────────
            PHASE_SIGNATURES = {
                "cleanup_drop":             "IF OBJECT_ID('tempdb..#LineBreakdown')",
                "breakdown_block":          "INTO #LineBreakdown",
                "final_join_detail_column": "[Price_Breakdown_Detail]",
                "auditor_queries":          "AUDITOR PRICE BREAKDOWN SUMMARY QUERIES",
            }
            signature = PHASE_SIGNATURES.get(phase)
            if signature and signature.upper() in current_sql.upper():
                print(f"  SKIP (idempotent): phase={phase} signature already present in SQL")
                append_log(file_outputs["log_file"], f"IDEMPOTENT_SKIP | {phase} | signature found")
                continue

            # Role-table proximity check (non-fatal warning)
            if role_table and role_table != "cleanup_block":
                if not validate_anchor_against_table(current_sql, anchor, role_table):
                    warn = (f"role_table '{role_table}' not in 3000-char window "
                            f"before anchor — possible drift")
                    print(f"  WARNING: {warn}")
                    append_log(file_outputs["log_file"],
                               f"ANCHOR_WARNING | {phase} | {warn}")

            # ── Generate content ──────────────────────────────────────────────
            context      = get_context(current_sql, anchor)
            content_msgs = build_content_messages(
                system_prompt, reference_sql, current_sql,
                context, analysis, ins,
                confirmed_lb_columns=confirmed_lb_columns,
            )

            ins_start = time.perf_counter()
            try:
                # Per-phase token limits.
                # breakdown_block generates #bRanked + #bSlots + full #LineBreakdown
                # SELECT with ServiceCategory/PricingMethod/LinePayment CASEs —
                # easily 600-800 lines, which saturates 16k tokens.
                # Other phases produce small outputs and 16k is sufficient.
                PHASE_MAX_TOKENS = {
                    "breakdown_block":          40000,
                    "auditor_queries":          8000,
                    "final_join_detail_column": 4000,
                    "cleanup_drop":             1000,
                }
                phase_max_tokens = PHASE_MAX_TOKENS.get(phase, 40000)

                rewrite, msg, elapsed, cost = call_claude_with_retry(
                    client, content_msgs, f"GENERATE {phase}",
                    max_tokens=phase_max_tokens,
                )

                inp, out, _, _ = usage_numbers(msg)
                file_total_input  += inp
                file_total_output += out
                file_total_time   += elapsed
                file_total_cost   += cost

                new_content = rewrite.get("new_content", "")
                if not new_content or not new_content.strip():
                    _skip(file_id, insertion_id, description, anchor, position,
                          "Empty new_content returned", file_outputs,
                          anchor_line=anchor_line)
                    failed += 1
                    continue

                # ── Post-generation validation for breakdown_block ────────
                if phase == "breakdown_block":

                    # Check 1: correct source table + no slot recomputation
                    source_issues = validate_breakdown_source(new_content, analysis)
                    if source_issues:
                        for si in source_issues:
                            code = si.split(" | ")[0]
                            print(f"  ✗ FATAL {code} — breakdown_block rejected: "
                                  f"{si.split(' | ', 1)[1][:120]}")
                            append_log(file_outputs["log_file"], f"FATAL | {si}")
                        _skip(file_id, insertion_id, description, anchor, position,
                              f"breakdown_block source validation failed: "
                              f"{[s.split(' | ')[0] for s in source_issues]}",
                              file_outputs, anchor_line=anchor_line)
                        failed += 1
                        continue

                    # Check 2: ContractEffectiveDateFrom/To present + period coverage
                    cy_issues = validate_contract_years(new_content, analysis)
                    if cy_issues:
                        for ci in cy_issues:
                            code = ci.split(' | ')[0]
                            print(f'  ✗ FATAL {code} — retrying with correction note')
                            append_log(file_outputs['log_file'], f'FATAL | {ci}')
                        if attempt_count < 2:
                            attempt_count += 1
                            print(f'  Retrying breakdown_block (contract year fix, attempt {attempt_count+1})…')
                            periods_json = analysis.get('contract_periods', [])
                            # Determine which specific errors fired so the retry note is accurate
                            cy_codes = [ci.split(' | ')[0] for ci in cy_issues]
                            has_missing_col = any(c in cy_codes for c in (
                                'MISSING_CONTRACT_DATE_FROM', 'MISSING_CONTRACT_DATE_TO'))
                            has_gap = 'CONTRACT_PERIOD_GAP' in cy_codes

                            if has_missing_col and not has_gap:
                                retry_reason = (
                                    f'CONTRACT DATE COLUMNS MISSING: '
                                    f'The previous attempt omitted ContractEffectiveDateFrom '
                                    f'and/or ContractEffectiveDateTo from #LineBreakdown entirely. '
                                    f'You MUST add BOTH columns immediately after [BundledByNCCI].'
                                )
                            elif has_gap and not has_missing_col:
                                retry_reason = (
                                    f'CONTRACT DATE FORMAT MISMATCH: '
                                    f'ContractEffectiveDateFrom/ContractEffectiveDateTo columns '
                                    f'were generated but one or more period date literals are wrong '
                                    f'or missing. You MUST use exact ISO-format date strings '
                                    f"(e.g. '2022-01-01') — NOT MM/DD/YYYY, NOT YYYYMMDD. "
                                    f'Every period must appear as a WHEN b.[ServiceDate] BETWEEN '
                                    f"'date_from' AND 'date_to' branch."
                                )
                            else:
                                retry_reason = (
                                    f'CONTRACT DATE COLUMNS MISSING OR MALFORMED: '
                                    f'Either ContractEffectiveDateFrom/ContractEffectiveDateTo '
                                    f'are absent, or the date literals do not match the expected '
                                    f'ISO format (YYYY-MM-DD). Both columns are MANDATORY.'
                                )

                            ins_cy_retry = dict(ins)
                            ins_cy_retry['_retry_note'] = (
                                f'CRITICAL RETRY — {retry_reason} '
                                f'Use EVERY period from contract_periods: {periods_json}. '
                                f'No gaps. ELSE NULL. These columns are NON-NEGOTIABLE.'
                            )
                            content_msgs_cy = build_content_messages(
                                system_prompt, reference_sql, current_sql,
                                context, analysis, ins_cy_retry,
                                confirmed_lb_columns=confirmed_lb_columns,
                            )
                            rewrite, msg, elapsed_cy, cost_cy = call_claude_with_retry(
                                client, content_msgs_cy, f'RETRY {phase} (contract years)',
                                max_tokens=phase_max_tokens,
                            )
                            inp_cy, out_cy, _, _ = usage_numbers(msg)
                            file_total_input  += inp_cy
                            file_total_output += out_cy
                            file_total_time   += elapsed_cy
                            file_total_cost   += cost_cy
                            new_content = rewrite.get('new_content', '')
                            cy_issues2 = validate_contract_years(new_content, analysis)
                            if cy_issues2:
                                print(f'  ✗ Contract date columns still missing after retry: '
                                      f'{[c.split(" | ")[0] for c in cy_issues2]}')
                                append_log(file_outputs['log_file'],
                                           f'CONTRACT_YEAR_FAIL_RETRY | {cy_issues2}')
                                _skip(file_id, insertion_id, description, anchor, position,
                                      f'Contract date columns missing after retry',
                                      file_outputs, anchor_line=anchor_line)
                                failed += 1
                                continue
                            else:
                                print('  ✓ Contract date columns present after retry')
                        else:
                            _skip(file_id, insertion_id, description, anchor, position,
                                  f'Contract date columns missing, retry limit reached',
                                  file_outputs, anchor_line=anchor_line)
                            failed += 1
                            continue

                    # Check 3: rn_ alias consistency
                    missing_rn = validate_branked_consistency(new_content)
                    if missing_rn:
                        print(f"  ⚠ rn_ consistency check FAILED — referenced but not defined: {missing_rn}")
                        append_log(file_outputs["log_file"],
                                   f"RN_CONSISTENCY_FAIL | {missing_rn}")
                        if attempt_count < 2:  # retry once
                            attempt_count += 1
                            print(f"  Retrying breakdown_block generation (attempt {attempt_count+1})…")
                            # Add the missing aliases to the instruction context
                            ins_retry = dict(ins)
                            ins_retry["_retry_note"] = (
                                f"CRITICAL RETRY: The previous attempt referenced these rn_ aliases "
                                f"in #LineBreakdown but never defined them in #bRanked: {missing_rn}. "
                                f"IMPORTANT: If any of these are 'rn_CAT', 'rn_ALIAS', 'rn_MAX_ALIAS', "
                                f"'rn_HYBRID_ALIAS', or any other template placeholder name from the "
                                f"instructions — those are NOT real aliases. Replace them with the actual "
                                f"rn_CategoryName aliases for the specific categories in this script "
                                f"(e.g. rn_ER, rn_Gamma_Knife, rn_Cardiac_Cath). "
                                f"Every rn_ alias referenced in #LineBreakdown MUST be defined in #bRanked "
                                f"with the exact same spelling. Add any missing ROW_NUMBER() lines now."
                            )
                            content_msgs_retry = build_content_messages(
                                system_prompt, reference_sql, current_sql,
                                context, analysis, ins_retry,
                                confirmed_lb_columns=confirmed_lb_columns,
                            )
                            rewrite, msg, elapsed2, cost2 = call_claude_with_retry(
                                client, content_msgs_retry, f"RETRY {phase}",
                                max_tokens=phase_max_tokens,
                            )
                            inp2, out2, _, _ = usage_numbers(msg)
                            file_total_input  += inp2
                            file_total_output += out2
                            file_total_time   += elapsed2
                            file_total_cost   += cost2
                            new_content = rewrite.get("new_content", "")
                            # Check again
                            missing_rn2 = validate_branked_consistency(new_content)
                            if missing_rn2:
                                print(f"  ⚠ rn_ still missing after retry: {missing_rn2}")
                                append_log(file_outputs["log_file"],
                                           f"RN_CONSISTENCY_FAIL_RETRY | {missing_rn2}")
                            else:
                                print(f"  ✓ rn_ consistency check passed after retry")

                # ── Apply ─────────────────────────────────────────────────────
                current_sql, success = apply_insertion(
                    current_sql, anchor, new_content, position
                )
                lines_added         = len(new_content.strip().splitlines())
                chars_added         = len(new_content)
                insertion_ends_line = anchor_line + lines_added if success else None

                if success:
                    successful += 1
                    print(f"  ✓ Inserted {lines_added} lines / {chars_added} chars")
                    print(f"    Lines {anchor_line} → {insertion_ends_line} in patched SQL")

                    # Extract confirmed #LineBreakdown columns after breakdown_block
                    # so downstream phases (final_join_detail_column, auditor_queries)
                    # know exactly which columns exist and can reference safely.
                    if phase == "breakdown_block":
                        confirmed_lb_columns = rewrite.get("confirmed_output_columns", [])
                        if confirmed_lb_columns:
                            print(f"  ✓ Confirmed #LineBreakdown columns: {confirmed_lb_columns}")
                            append_log(file_outputs["log_file"],
                                       f"LB_COLUMNS | {confirmed_lb_columns}")
                        else:
                            print("  WARNING: breakdown_block returned no confirmed_output_columns")
                            append_log(file_outputs["log_file"],
                                       "LB_COLUMNS_MISSING | breakdown_block did not return confirmed_output_columns")

                    insertions_log["insertions"].append({
                        "insertion_id":          insertion_id,
                        "phase":                 phase,
                        "description":           description,
                        "role_table":            role_table,
                        "anchor_preview":        anchor[:150],
                        "anchor_line_number":    anchor_line,
                        "insertion_ends_at_line": insertion_ends_line,
                        "lines_added":           lines_added,
                        "chars_added":           chars_added,
                        "new_content_preview":   new_content[:600],
                        **({"confirmed_output_columns": confirmed_lb_columns}
                           if phase == "breakdown_block" else {}),
                    })
                else:
                    failed += 1
                    print(f"  ✗ apply_insertion returned False for {phase}")

                _log_db_insertion(
                    file_id, insertion_id, description,
                    anchor, position,
                    anchor_line, insertion_ends_line,
                    success, new_content,
                    lines_added if success else 0,
                    chars_added if success else 0,
                    elapsed,
                    None if success else "apply_insertion returned False",
                )

                append_log(
                    file_outputs["log_file"],
                    f"INSERTION {insertion_id} | phase={phase} "
                    f"anchor_line={anchor_line} ends_line={insertion_ends_line} "
                    f"in={inp} out={out} time={elapsed:.2f} "
                    f"cost={cost:.4f} success={success} lines={lines_added}",
                )

            except Exception as e:
                print(f"  ERROR: {e}")
                failed += 1
                _log_db_insertion(
                    file_id, insertion_id, description,
                    anchor, position,
                    anchor_line, None,
                    False, "", 0, 0,
                    time.perf_counter() - ins_start, str(e),
                )

        # ── Save outputs ──────────────────────────────────────────────────────
        file_outputs["patched_sql"].write_text(current_sql, encoding="utf-8")
        _save_log(file_outputs, insertions_log, input_file_path,
                  len(insertions), successful, failed,
                  file_total_input, file_total_output,
                  file_total_time, file_total_cost)

        file_runtime = time.perf_counter() - file_start

        if failed == 0 and not missing_phases:
            status = "SUCCESS"
        elif successful > 0:
            status = "PARTIAL"
        else:
            status = "FAILED"

        _end_db_file(file_id, status, file_outputs,
                     len(insertions), successful, file_total_cost, file_runtime)

        print(f"\n  Result   : {successful}/{len(insertions)} successful | "
              f"${file_total_cost:.4f} | {file_runtime:.1f}s")
        if missing_phases:
            print(f"  ⚠ Missing: {missing_phases}")
        print(f"  Output   : {file_outputs['patched_sql']}")
        print(f"  Analysis : {file_outputs['analysis_log']}")

        return _result(input_file_path.name, status == "SUCCESS",
                       len(insertions), successful,
                       file_total_cost, file_runtime,
                       missing_phases=missing_phases)

    except Exception as e:
        print(f"  FATAL: {e}")
        append_log(file_outputs["log_file"], f"ERROR | {input_file_path.name} | {e}")
        file_runtime = time.perf_counter() - file_start
        _end_db_file(file_id, "FAILED", file_outputs,
                     0, 0, 0.0, file_runtime, str(e))
        return _result(input_file_path.name, False, 0, 0,
                       0.0, file_runtime, error=str(e))


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _skip(file_id, ins_id, desc, anchor, position, reason, file_outputs,
          preview=None, anchor_line=None):
    print(f"  SKIP: {reason}")
    if preview:
        print(f"  Anchor preview: {preview}")
    append_log(file_outputs["log_file"],
               f"SKIP | insertion={ins_id} | {reason}")
    _log_db_insertion(file_id, ins_id, desc, anchor, position,
                      anchor_line, None,
                      False, "", 0, 0, 0.0, reason)


def _save_log(file_outputs, insertions_log, input_file_path,
              total, successful, failed, inp, out, t, cost):
    with open(file_outputs["insertions_log"], "w", encoding="utf-8") as f:
        json.dump(insertions_log, f, indent=2)
    append_log(
        file_outputs["log_file"],
        f"FILE_SUMMARY | {input_file_path.name} | "
        f"insertions={total} success={successful} failed={failed} "
        f"input={inp} output={out} time={t:.2f} cost={cost:.4f}",
    )


def _end_db_file(file_id, status, file_outputs, found, applied,
                 cost, secs, error=None):
    global db_manager
    if db_manager and file_id:
        try:
            db_manager.end_file(
                file_id=file_id,
                status=status,
                output_file_path=str(file_outputs["patched_sql"]),
                edits_found=found,
                edits_applied=applied,
                cost=cost,
                processing_seconds=secs,
                **({"error_message": error} if error else {}),
            )
        except Exception as e:
            print(f"  DB error ending file: {e}")


def _log_db_insertion(file_id, ins_id, desc,
                      anchor, position,
                      anchor_line, insertion_ends_line,
                      success, new_content,
                      lines, chars, secs, error):
    """
    Log insertion to database.

    DB schema required:
      position               NVARCHAR(500)  — anchor text (not integer offset)
      anchor_line_number     INT NULL       — line where anchor ends in source SQL
      insertion_ends_at_line INT NULL       — anchor_line + lines_added in patched SQL

    Run once before first use:
      ALTER TABLE dbo.SQL_Processing_Edits
          ALTER COLUMN position NVARCHAR(500);
      ALTER TABLE dbo.SQL_Processing_Edits
          ADD anchor_line_number    INT NULL,
              insertion_ends_at_line INT NULL;
    """
    global db_manager
    if db_manager and file_id:
        try:
            # Truncate description to fit database column constraints
            # Most SQL databases have limited varchar sizes for description fields
            max_desc_length = 255  # Conservative limit for description column
            truncated_desc = desc[:max_desc_length] if desc else ""
            if len(desc) > max_desc_length:
                truncated_desc = desc[:max_desc_length-3] + "..."
                print(f"  Description truncated: {len(desc)} -> {len(truncated_desc)} chars")

            db_manager.log_edit(
                file_id=file_id,
                edit_number=ins_id,
                description=truncated_desc,          # Truncated to fit column
                position=anchor[:500],               # NVARCHAR(500) — anchor text
                match_type=position,
                success=success,
                old_snippet_full_text=anchor,
                new_snippet_full_text=new_content,
                lines_added=lines,
                characters_added=chars,
                processing_seconds=secs,
                error_message=error,
                anchor_line_number=anchor_line,
                insertion_ends_at_line=insertion_ends_line,
            )
        except Exception as e:
            print(f"  DB error logging insertion: {e}")


def _result(filename, success, found, applied, cost, t,
            error=None, missing_phases=None):
    r = {
        "filename":           filename,
        "success":            success,
        "insertions_found":   found,
        "insertions_applied": applied,
        "cost":               cost,
        "time":               t,
        "missing_phases":     missing_phases or [],
    }
    if error:
        r["error"] = error
    return r


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    global session_metrics, db_manager, total_cached_tokens, total_uncached_tokens

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("Set ANTHROPIC_API_KEY in .env")

    client = anthropic.Anthropic(api_key=api_key)
    session_metrics["start_time"] = time.perf_counter()

    print("Claude SQL Batch INSERT Pipeline  v9")
    print("Strategy: INDICATOR_FLAG support, rn_ naming enforcement, rn_ consistency retry")

    try:
        db_manager = DatabaseManager()
        ok = db_manager.test_connection()
        print(f"Database: {'OK' if ok else 'FAILED — no DB logging'}")
        if not ok:
            db_manager = None
    except Exception as e:
        print(f"Database unavailable ({e}) — no DB logging")
        db_manager = None

    INPUT_DIR.mkdir(exist_ok=True)
    BATCH_OUTPUT_DIR.mkdir(exist_ok=True)

    sql_files = list(INPUT_DIR.glob("*.sql"))
    if not sql_files:
        print(f"No .sql files in {INPUT_DIR}")
        return

    print(f"\nFiles to process: {len(sql_files)}")
    for f in sql_files:
        print(f"  {f.name}")

    print("\nLoading system prompt + reference SQL…")
    system_prompt = read_docx(SYSTEM_PROMPT_PATH)
    reference_sql = read_text(REFERENCE_SQL_PATH)

    session_id = None
    if db_manager:
        try:
            session_id = db_manager.start_session(
                script_version="sql_insert_pipeline_v9.py",
                reference_file=REFERENCE_SQL_PATH.name,
                approach="Role-based anchor insertion v9 — INDICATOR_FLAG + rn_ naming + consistency retry",
            )
            print(f"DB session: {session_id}")
        except Exception as e:
            print(f"DB session start error: {e}")

    results = []
    for sql_file in sql_files:
        file_outputs = create_file_output_structure(sql_file.name, BATCH_OUTPUT_DIR)
        result       = process_single_file(
            client, sql_file, system_prompt, reference_sql,
            file_outputs, session_id,
        )
        results.append(result)

        session_metrics["files_processed"] += 1
        if result["success"]:
            session_metrics["files_successful"]            += 1
            session_metrics["total_insertions"]            += result["insertions_found"]
            session_metrics["total_successful_insertions"] += result["insertions_applied"]
        else:
            session_metrics["files_failed"] += 1
        session_metrics["total_cost"] += result["cost"]

    total_time      = time.perf_counter() - session_metrics["start_time"]
    all_tokens      = total_uncached_tokens + total_cached_tokens
    cost_no_cache   = (all_tokens / 1_000_000) * INPUT_COST_PER_MILLION
    cost_with_cache = session_metrics["total_cost"]
    savings         = cost_no_cache - cost_with_cache
    savings_pct     = (savings / cost_no_cache * 100) if cost_no_cache > 0 else 0

    if db_manager and session_id:
        try:
            db_manager.end_session(
                session_id=session_id,
                status="COMPLETED" if session_metrics["files_failed"] == 0 else "PARTIAL",
                total_files=session_metrics["files_processed"],
                files_successful=session_metrics["files_successful"],
                files_failed=session_metrics["files_failed"],
                total_edits=session_metrics["total_insertions"],
                total_successful_edits=session_metrics["total_successful_insertions"],
                total_cost=cost_with_cache,
                total_savings=savings,
                processing_seconds=total_time,
            )
            db_manager.disconnect()
        except Exception as e:
            print(f"DB session end error: {e}")

    print(f"\n{'='*60}")
    print("BATCH SUMMARY")
    print(f"  Files      : {session_metrics['files_successful']}/{session_metrics['files_processed']} successful")
    print(f"  Insertions : {session_metrics['total_successful_insertions']}/{session_metrics['total_insertions']} applied")
    print(f"  Cost       : ${cost_with_cache:.4f}  (saved ${savings:.4f} / {savings_pct:.0f}% via cache)")
    print(f"  Time       : {total_time:.1f}s")
    print()
    for r in results:
        icon = "✓" if r["success"] else "✗"
        missing_note = f" ⚠ missing={r['missing_phases']}" if r.get("missing_phases") else ""
        print(f"  {icon} {r['filename']} | "
              f"{r['insertions_applied']}/{r['insertions_found']} | "
              f"${r['cost']:.4f} | {r['time']:.1f}s{missing_note}")
        if not r["success"] and r.get("error"):
            print(f"    Error: {r['error']}")

    summary_log = BATCH_OUTPUT_DIR / "batch_summary.log"
    with open(summary_log, "w", encoding="utf-8") as f:
        f.write(
            f"session_id={session_id} "
            f"files={session_metrics['files_processed']} "
            f"success={session_metrics['files_successful']} "
            f"failed={session_metrics['files_failed']} "
            f"insertions={session_metrics['total_insertions']} "
            f"applied={session_metrics['total_successful_insertions']} "
            f"cost={cost_with_cache:.4f} savings={savings:.4f} "
            f"time={total_time:.2f}\n\nFILE RESULTS:\n"
        )
        for r in results:
            f.write(
                f"{r['filename']} | success={r['success']} | "
                f"insertions={r['insertions_applied']}/{r['insertions_found']} | "
                f"missing={r['missing_phases']} | "
                f"cost=${r['cost']:.4f}\n"
            )

    print(f"\nOutputs → {BATCH_OUTPUT_DIR}")


if __name__ == "__main__":
    main()