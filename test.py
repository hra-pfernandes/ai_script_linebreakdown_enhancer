import os
import time
from pathlib import Path
from typing import Optional

from dotenv import load_dotenv
import anthropic
from docx import Document

# Load environment variables from .env file
load_dotenv()


# ==============================
# CONFIG
# ==============================

MODEL = "claude-sonnet-4-6"
MAX_OUTPUT_TOKENS = 64000
MAX_CONTINUATIONS = 10

INPUT_COST_PER_MILLION = 3.00
OUTPUT_COST_PER_MILLION = 15.00

BASE_DIR = Path(__file__).resolve().parent

SYSTEM_PROMPT_FILE = BASE_DIR / "data/HRA_Script_Line_Level_Breakdown_System_Prompt_v2.docx"
REFERENCE_SQL_FILE = BASE_DIR / "data/AI_NYP_COL_AETNA_FINAL2.sql"
TARGET_SQL_FILE = BASE_DIR / "data/NYP_COL_Aetna_Commercial_OP_Complete 1.sql"

OUTPUT_SQL_FILE = BASE_DIR / "output/result.sql"
METRICS_LOG = BASE_DIR / "output/run_metrics.log"


# ==============================
# HELPERS
# ==============================

def read_docx(path: Path) -> str:
    """
    Read a Word document and extract text content.
    
    Args:
        path: Path to the .docx file
        
    Returns:
        Extracted text content
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        Exception: If the file is corrupted or unreadable
    """
    if not path.exists():
        raise FileNotFoundError(f"System prompt file not found: {path}")
    
    try:
        doc = Document(str(path))
        parts = []

        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)
    except Exception as e:
        raise Exception(f"Failed to read document {path}: {e}")


def read_file(path: Path) -> str:
    """
    Read a text file with error handling.
    
    Args:
        path: Path to the file
        
    Returns:
        File content as string
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        Exception: If the file cannot be read
    """
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        raise Exception(f"Failed to read file {path}: {e}")


def estimate_cost(input_tokens, output_tokens):
    input_cost = (input_tokens / 1_000_000) * INPUT_COST_PER_MILLION
    output_cost = (output_tokens / 1_000_000) * OUTPUT_COST_PER_MILLION
    return input_cost + output_cost


def merge_sql(existing, new):
    """
    Merge continuation SQL safely by removing overlap
    """
    max_overlap = 4000
    for i in range(max_overlap, 0, -1):
        if existing.endswith(new[:i]):
            return existing + new[i:]
    return existing + new


def log_metrics(text):
    METRICS_LOG.parent.mkdir(parents=True, exist_ok=True)
    with open(METRICS_LOG, "a", encoding="utf-8") as f:
        f.write(text + "\n")


# ==============================
# CLAUDE CALL
# ==============================

def call_claude(client, system_prompt, user_prompt):

    parts = []
    start = time.time()

    with client.messages.stream(
        model=MODEL,
        max_tokens=MAX_OUTPUT_TOKENS,
        system=system_prompt,
        messages=[{"role": "user", "content": user_prompt}],
    ) as stream:

        for text in stream.text_stream:
            print(text, end="", flush=True)
            parts.append(text)

        message = stream.get_final_message()

    duration = time.time() - start
    return "".join(parts), message, duration


# ==============================
# MAIN
# ==============================

def main():

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("Set ANTHROPIC_API_KEY")

    client = anthropic.Anthropic(api_key=api_key)

    system_prompt = read_docx(SYSTEM_PROMPT_FILE)
    reference_sql = read_file(REFERENCE_SQL_FILE)
    target_sql = read_file(TARGET_SQL_FILE)

    prompt = f"""
Modify the TARGET SQL using the REFERENCE SQL as pattern.

Return ONLY the final SQL.

REFERENCE SQL:
{reference_sql}

TARGET SQL:
{target_sql}
"""

    total_input = 0
    total_output = 0
    total_time = 0
    sql_result = ""

    start_total = time.time()

    for i in range(1, MAX_CONTINUATIONS + 1):

        print(f"\n\n===== CLAUDE PASS {i} =====\n")

        output, message, duration = call_claude(client, system_prompt, prompt)

        sql_result = merge_sql(sql_result, output)

        usage = message.usage
        input_tokens = usage.input_tokens
        output_tokens = usage.output_tokens
        stop_reason = message.stop_reason

        total_input += input_tokens
        total_output += output_tokens
        total_time += duration

        cost = estimate_cost(input_tokens, output_tokens)

        print("\n\nPASS METRICS")
        print("Input tokens:", input_tokens)
        print("Output tokens:", output_tokens)
        print("Stop reason:", stop_reason)
        print("Time:", round(duration, 2), "sec")
        print("Cost:", round(cost, 4))

        log_metrics(
            f"pass={i} input={input_tokens} output={output_tokens} time={duration:.2f} cost={cost:.4f}"
        )

        if stop_reason != "max_tokens":
            break

        print("\n--- CONTINUING GENERATION ---")

        prompt = f"""
Continue the SQL exactly where it stopped.

Return only the continuation.

SQL so far (end section):

{sql_result[-10000:]}
"""

    total_runtime = time.time() - start_total
    total_cost = estimate_cost(total_input, total_output)

    OUTPUT_SQL_FILE.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_SQL_FILE.write_text(sql_result, encoding="utf-8")

    print("\n\n==============================")
    print("SESSION COMPLETE")
    print("==============================")

    print("Total input tokens:", total_input)
    print("Total output tokens:", total_output)
    print("Total time:", round(total_runtime, 2), "sec")
    print("Total cost:", round(total_cost, 4))
    print("Output file:", OUTPUT_SQL_FILE)

    log_metrics(
        f"SESSION total_input={total_input} total_output={total_output} "
        f"time={total_runtime:.2f} cost={total_cost:.4f}"
    )


if __name__ == "__main__":
    main()