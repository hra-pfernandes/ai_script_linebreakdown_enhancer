# Claude SQL Batch INSERT Pipeline  v9
# ─────────────────────────────────────────────────────────────────────────────
# Strategy   : Pure additive insertions — no replacement of existing SQL.
#
# Four REQUIRED insertion phases (minimum — more allowed per script):
#   1. cleanup_drop             → TOP of script, cleanup block
#   2. breakdown_block          → after HIERARCHY, before TOTAL_PRICE
#   3. final_join_detail_column → inside FINAL_JOIN after ExpectedDetailed
#   4. auditor_queries          → END of script, after final OUTPUT SELECT
#
# Anchoring  : Discovery derives anchors from FUNCTIONAL ROLE identification
#              (Phase 1 of system prompt), not arbitrary text snippets.
#
# Line nums  : anchor_line_number and insertion_ends_at_line are captured for
#              every insertion — making patched SQL easy to navigate/review.
#
# DB logging : position column is NVARCHAR(500) — stores anchor text.
#              anchor_line_number and insertion_ends_at_line are INT columns.
#              Run the migration script below before first use:
#
#   ALTER TABLE dbo.SQL_Processing_Edits
#       ALTER COLUMN position NVARCHAR(500);
#   ALTER TABLE dbo.SQL_Processing_Edits
#       ADD anchor_line_number    INT NULL,
#           insertion_ends_at_line INT NULL;
# ─────────────────────────────────────────────────────────────────────────────

import json
import os
import re
import random
import time
from datetime import datetime
from pathlib import Path

import anthropic
from database_manager import DatabaseManager
from docx import Document
from dotenv import load_dotenv

load_dotenv()

MODEL = "claude-sonnet-4-6"

INPUT_COST_PER_MILLION      = 3.00
OUTPUT_COST_PER_MILLION     = 15.00
CACHE_READ_COST_PER_MILLION = 0.15

BASE_DIR           = Path(__file__).resolve().parent
SYSTEM_PROMPT_PATH = BASE_DIR / "data" / "HRA_Script_Line_Level_Breakdown_System_Prompt_v8.docx"
REFERENCE_SQL_PATH = BASE_DIR / "data" / "NYP_COL_Aetna_Commercial_OP_Complete 1_final.sql"
INPUT_DIR          = BASE_DIR / "input"
BATCH_OUTPUT_DIR   = BASE_DIR / "batch_output"

CONTEXT_CHARS_BEFORE = 1000
CONTEXT_CHARS_AFTER  = 1000

# Every script must have at least these four phases.
REQUIRED_PHASES = {
    "cleanup_drop",
    "breakdown_block",
    "final_join_detail_column",
    "auditor_queries",
}

total_cached_tokens   = 0
total_uncached_tokens = 0

session_metrics = {
    "start_time": None,
    "files_processed": 0,
    "files_successful": 0,
    "files_failed": 0,
    "total_insertions": 0,
    "total_successful_insertions": 0,
    "total_cost": 0.0,
    "total_savings": 0.0,
}

db_manager = None


# ─────────────────────────────────────────────────────────────────────────────
# I/O helpers
# ─────────────────────────────────────────────────────────────────────────────

def read_docx(path):
    doc   = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for r in t.rows:
            cells = [c.text.strip() for c in r.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_text(path):
    return path.read_text(encoding="utf-8", errors="ignore")


def append_log(log_path, line):
    log_path.parent.mkdir(parents=True, exist_ok=True)
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(line + "\n")


def create_file_output_structure(input_filename, batch_output_dir):
    base_name       = Path(input_filename).stem
    file_output_dir = batch_output_dir / base_name
    file_output_dir.mkdir(parents=True, exist_ok=True)
    return {
        "output_dir":     file_output_dir,
        "patched_sql":    file_output_dir / f"{base_name}_patched.sql",
        "log_file":       file_output_dir / f"{base_name}_metrics.log",
        "insertions_log": file_output_dir / f"{base_name}_insertions.json",
        "analysis_log":   file_output_dir / f"{base_name}_analysis.json",
    }


# ─────────────────────────────────────────────────────────────────────────────
# Cost helpers
# ─────────────────────────────────────────────────────────────────────────────

def estimate_cost(input_tokens, output_tokens, cached_tokens):
    return (
          (input_tokens  / 1_000_000) * INPUT_COST_PER_MILLION
        + (cached_tokens / 1_000_000) * CACHE_READ_COST_PER_MILLION
        + (output_tokens / 1_000_000) * OUTPUT_COST_PER_MILLION
    )


def usage_numbers(msg):
    u = getattr(msg, "usage", None)
    if not u:
        return 0, 0, 0, 0
    return (
        getattr(u, "input_tokens", 0),
        getattr(u, "output_tokens", 0),
        getattr(u, "cache_creation_input_tokens", 0),
        getattr(u, "cache_read_input_tokens", 0),
    )


# ─────────────────────────────────────────────────────────────────────────────
# JSON extraction
# ─────────────────────────────────────────────────────────────────────────────

def extract_json(text):
    """
    Extract and parse a JSON object from Claude's response.

    Handles three common response formats:
      1. Raw JSON   — Claude returns { ... } directly
      2. Code fence — Claude wraps in ```json ... ``` or ``` ... ```
      3. Mixed text — JSON embedded somewhere in prose

    Enhanced with control character cleaning for SQL code in JSON strings.
    """
    if not text or not text.strip():
        return {"error": "Empty or whitespace-only response", "raw_text": text}
    
    text = text.strip()
    original_length = len(text)

    def clean_json_for_parsing(json_text):
        """Clean control characters that break JSON parsing while preserving structure."""
        # Replace problematic control characters in string values
        # This regex finds string values and replaces control chars within them
        import re
        
        def replace_in_string(match):
            string_content = match.group(0)
            # Replace literal newlines, tabs, etc. with escaped versions
            cleaned = string_content.replace('\\n', '\\n')  # Already escaped - keep as is
            cleaned = cleaned.replace('\n', '\\n')          # Literal newline - escape it
            cleaned = cleaned.replace('\r', '\\r')          # Literal carriage return
            cleaned = cleaned.replace('\t', '\\t')          # Literal tab
            cleaned = cleaned.replace('\b', '\\b')          # Literal backspace
            cleaned = cleaned.replace('\f', '\\f')          # Literal form feed
            return cleaned
        
        # Find all string values (including multiline ones) and clean them
        # This regex matches: "..." including escaped quotes and newlines
        pattern = r'"(?:[^"\\]|\\.)*"'
        return re.sub(pattern, replace_in_string, json_text, flags=re.DOTALL)

    # Strategy 1: strip markdown code fences if present, then parse
    fence_stripped = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    fence_stripped = re.sub(r"\s*```\s*$", "", fence_stripped).strip()
    
    if fence_stripped != text:
        print(f"  Stripped code fences: {len(text)} -> {len(fence_stripped)} chars")
    
    # Clean control characters before parsing
    cleaned_text = clean_json_for_parsing(fence_stripped)
    if cleaned_text != fence_stripped:
        print(f"  Cleaned control characters: {len(fence_stripped)} -> {len(cleaned_text)} chars")
    
    try:
        parsed = json.loads(cleaned_text)
        print(f"  JSON parsed successfully (fence-stripped + cleaned, {len(cleaned_text)} chars)")
        return parsed
    except json.JSONDecodeError as e:
        print(f"  Fence-stripped + cleaned parse failed: {e}")

    # Strategy 2: direct parse (no fences, but still clean)
    cleaned_direct = clean_json_for_parsing(text)
    try:
        parsed = json.loads(cleaned_direct)
        print(f"  JSON parsed successfully (direct + cleaned, {len(cleaned_direct)} chars)")
        return parsed
    except json.JSONDecodeError as e:
        print(f"  Direct + cleaned parse failed: {e}")

    # Strategy 3: extract outermost { ... } using rfind for the closing brace
    start = text.find("{")
    end   = text.rfind("}")
    
    if start == -1:
        print(f"  No opening brace found in {len(text)} chars")
        return {"error": "No JSON opening brace found", "raw_text": text[:500]}
    
    if end == -1:
        print(f"  No closing brace found in {len(text)} chars")
        return {"error": "No JSON closing brace found - possibly truncated", "raw_text": text[:500]}
    
    if end <= start:
        print(f"  Invalid brace positions: start={start}, end={end}")
        return {"error": "Invalid JSON brace positions", "raw_text": text[:500]}
    
    json_candidate = text[start : end + 1]
    print(f"  Extracting JSON candidate: {len(json_candidate)} chars from positions {start}-{end}")
    
    # Clean the extracted candidate
    cleaned_candidate = clean_json_for_parsing(json_candidate)
    
    try:
        parsed = json.loads(cleaned_candidate)
        print(f"  JSON parsed successfully (extracted + cleaned, {len(cleaned_candidate)} chars)")
        return parsed
    except json.JSONDecodeError as e:
        print(f"  Extracted + cleaned JSON parse failed: {e}")
        print(f"  JSON candidate preview: {json_candidate[:200]}...")
        
        # Check for common truncation indicators
        if json_candidate.endswith("...") or "..." in json_candidate:
            return {"error": "JSON appears truncated (contains ellipsis)", "raw_text": text[:500]}

    # All strategies failed - provide detailed error info
    error_info = {
        "error": "Could not parse JSON with any strategy",
        "original_length": original_length,
        "cleaned_length": len(cleaned_text),
        "has_opening_brace": start != -1,
        "has_closing_brace": end != -1,
        "brace_positions": f"{start}-{end}" if start != -1 and end != -1 else "N/A",
        "contains_ellipsis": "..." in text,
        "raw_text": text[:500]
    }
    
    print(f"  JSON parse error. Length: {original_length}, Braces: {start}-{end}, Preview: {text[:100]}...")
    return error_info


# ─────────────────────────────────────────────────────────────────────────────
# Line number helper
# ─────────────────────────────────────────────────────────────────────────────

def get_line_number(sql: str, char_pos: int) -> int:
    """Return 1-based line number for a character position in sql."""
    return sql[:char_pos].count("\n") + 1


def validate_branked_consistency(new_content: str) -> list[str]:
    """
    After breakdown_block generation, verify that every rn_ alias referenced
    in #LineBreakdown is also defined in #bRanked within the same new_content.
    Returns a list of missing aliases (empty = all good).
    """
    import re
    # Find all rn_ aliases DEFINED in #bRanked (AS rn_xxx pattern)
    defined = set(re.findall(r'\bAS\s+(rn_\w+)', new_content, re.IGNORECASE))
    # Find all rn_ aliases REFERENCED anywhere after #bRanked (b.rn_xxx or rn_xxx =)
    referenced = set(re.findall(r'\b(rn_\w+)\b', new_content, re.IGNORECASE))
    # Missing = referenced but not defined
    missing = [r for r in referenced if r not in defined]
    return missing


# ─────────────────────────────────────────────────────────────────────────────
# Anchor helpers
# ─────────────────────────────────────────────────────────────────────────────

def find_anchor(sql: str, anchor: str) -> int | None:
    """
    Return the character position immediately AFTER `anchor` in `sql`.
    Tries exact match first, then whitespace-normalised match.
    """
    # 1. Exact
    pos = sql.find(anchor)
    if pos != -1:
        return pos + len(anchor)

    # 2. Normalised
    sql_norm    = re.sub(r"\s+", " ", sql)
    anchor_norm = re.sub(r"\s+", " ", anchor.strip())
    pos_norm    = sql_norm.find(anchor_norm)
    if pos_norm != -1:
        target_non_ws = len(re.sub(r"\s+", "", sql_norm[: pos_norm + len(anchor_norm)]))
        seen = 0
        for i, ch in enumerate(sql):
            if not ch.isspace():
                seen += 1
            if seen >= target_non_ws:
                return i + 1

    return None


def validate_anchor_against_table(sql: str, anchor: str, table_name: str) -> bool:
    """
    Confirm table_name appears anywhere in the SQL preceding the anchor.
    Searches the full preceding text — a fixed window would fail after large
    insertions shift the anchor far from its role table's definition.
    """
    anchor_pos = find_anchor(sql, anchor)
    if anchor_pos is None:
        return False
    preceding = sql[:anchor_pos]
    return table_name.upper() in preceding.upper()


def get_context(sql: str, anchor: str) -> str:
    pos = find_anchor(sql, anchor)
    if pos is None:
        return sql[:3000]
    return sql[max(0, pos - CONTEXT_CHARS_BEFORE) : min(len(sql), pos + CONTEXT_CHARS_AFTER)]


def apply_insertion(sql: str, anchor: str, new_content: str,
                    position: str) -> tuple[str, bool]:
    """Insert new_content immediately after (or before) the anchor."""
    anchor_end = find_anchor(sql, anchor)
    if anchor_end is None:
        print(f"  ERROR: anchor not found — '{anchor[:80]}'")
        return sql, False
    insert_at = anchor_end if position == "after" else max(0, anchor_end - len(anchor))
    return sql[:insert_at] + "\n\n" + new_content.strip() + "\n\n" + sql[insert_at:], True


# ─────────────────────────────────────────────────────────────────────────────
# Discovery prompt
# ─────────────────────────────────────────────────────────────────────────────

DISCOVERY_SCHEMA = """{
  "analysis": {
    "line_level_source_table":     "#ActualName",
    "hierarchy_table":             "#ActualName",
    "total_price_table":           "#ActualName",
    "final_join_table":            "#ActualName",
    "output_table":                "#ActualName",
    "has_ncci_bundle":             true,
    "has_window_reduction":        false,
    "window_reduction_categories": [],
    "has_aggregate_date_step":     false,
    "hybrid_categories":           [],
    "max_categories":              [],
    "sum_categories":              [],
    "indicator_flag_categories":   []
  },
  "insertions": [
    {
      "insertion_id":     1,
      "phase":            "cleanup_drop",
      "description":      "DROP TABLE safety lines for #bRanked, #bSlots, #LineBreakdown in the cleanup block",
      "role_anchor_table": "cleanup_block",
      "anchor_snippet":   "exact last ~80 chars of the last existing DROP TABLE line in the cleanup block at the TOP of the script",
      "insert_position":  "after",
      "occurrence_index": 1
    },
    {
      "insertion_id":     2,
      "phase":            "breakdown_block",
      "description":      "#bRanked temp table + #bSlots temp table (if window reduction) + SELECT INTO #LineBreakdown",
      "role_anchor_table": "#HierarchyTableActualName",
      "anchor_snippet":   "exact last ~80 chars of the final statement of the HIERARCHY step",
      "insert_position":  "after",
      "occurrence_index": 1
    },
    {
      "insertion_id":     3,
      "phase":            "final_join_detail_column",
      "description":      "Price_Breakdown_Detail correlated subquery as the last column in FINAL_JOIN SELECT",
      "role_anchor_table": "#FinalJoinTableActualName",
      "anchor_snippet":   "exact text of the INTO #TableName line of the FINAL_JOIN SELECT (e.g. 'Into #Step7') — this is the most reliable anchor because it is always unique",
      "insert_position":  "before",
      "occurrence_index": 1
    },
    {
      "insertion_id":     4,
      "phase":            "auditor_queries",
      "description":      "Auditor Query 1 (line detail + INSERT INTO persistent table) + Query 2 (reconciliation) appended after final OUTPUT SELECT",
      "role_anchor_table": "#OutputTableActualName",
      "anchor_snippet":   "exact last ~80 chars of the final OUTPUT SELECT statement",
      "insert_position":  "after",
      "occurrence_index": 1
    }
  ]
}"""


def build_discovery_messages(system_prompt, reference_sql, target_sql):
    content_blocks = [
        {
            "type": "text",
            "text": system_prompt,
            "cache_control": {"type": "ephemeral"},
        },
        {
            "type": "text",
            "text": f"REFERENCE SQL:\n{reference_sql}",
            "cache_control": {"type": "ephemeral"},
        },
        {
            "type": "text",
            "text": f"""CRITICAL: RETURN ONLY VALID JSON — NO PROSE, NO MARKDOWN, NO EXPLANATION.

════════════════════════════════════════════════════════════════
TASK
════════════════════════════════════════════════════════════════
Process the TARGET SQL below.

STEP A — Full Phase 1 analysis (system prompt):
  1A  Map every temp table to its functional role:
      FILTER, CHARGE_DETAIL, LINE_PRICING, NCCI_BUNDLE, AGGREGATE_DATE,
      AGGREGATE_ENC, HIERARCHY, TOTAL_PRICE, FINAL_JOIN, OUTPUT.
  1B  Extract all service categories from LINE_PRICING. Also identify any
      column whose CASE expression produces ONLY binary 0 or 1 values and
      never a dollar amount — these are INDICATOR_FLAG columns (e.g. COVID
      flag, CPT presence flags). List them separately as indicator_flag_categories.
  1C  Classify MAX vs SUM vs HYBRID from BOTH AGGREGATE_DATE (if present) AND
      AGGREGATE_ENC. A category that is MAX in AGGREGATE_DATE (GROUP BY
      EncounterID+ServiceDate) but SUM in AGGREGATE_ENC (GROUP BY EncounterID)
      is HYBRID — it requires PARTITION BY [EncounterID],[ServiceDate] in
      #bRanked. INDICATOR_FLAG categories are MAX in aggregation but contribute
      $0 to LinePayment — they still require ROW_NUMBER() in #bRanked and
      ServiceCategory labeling. Produce a numbered checklist of every MAX,
      HYBRID, and INDICATOR_FLAG category — this is the authoritative count
      for the #bRanked completeness check.
  1D  Extract the suppression hierarchy from HIERARCHY.
  1E  Identify WINDOW_REDUCTION categories (LEAD slot patterns).
  1F  Determine insertion points per system prompt rules.

STEP B — Return the JSON plan.

════════════════════════════════════════════════════════════════
REQUIRED INSERTIONS — you MUST include ALL FOUR phases.
Add extra insertions beyond these four if the script warrants it.
════════════════════════════════════════════════════════════════

  1. cleanup_drop
     WHERE: the cleanup / DROP TABLE block at the very TOP of the script
     WHAT:  a DROP TABLE safety line for #LineBreakdown

  2. breakdown_block
     WHERE: immediately after the HIERARCHY step, before TOTAL_PRICE
     WHAT:  SELECT INTO #bRanked (ranked source rows), optional SELECT INTO #bSlots
            (slot pivot for window reduction categories), SELECT INTO #LineBreakdown

  3. final_join_detail_column
     WHERE: immediately before the INTO clause of the FINAL_JOIN SELECT
            (the INTO #TableName line is the anchor — insert BEFORE it)
     WHAT:  the Price_Breakdown_Detail correlated subquery column as the
            last column in the SELECT, placed just before INTO

  4. auditor_queries
     WHERE: after the very last SELECT in the script (the OUTPUT step)
     WHAT:  Phase 4 auditor block — Query 1 (line detail SELECT) +
            INSERT INTO [Automation].[dbo].[LineBreakdown_Results] +
            Query 2 (reconciliation against TOTAL_PRICE table)

If ANY of these four phases is absent, your response will be rejected.

════════════════════════════════════════════════════════════════
ANCHOR RULES
════════════════════════════════════════════════════════════════
An anchor_snippet MUST be:
  • The FINAL statement (last ~80 chars) of the role-identified block.
  • Unique within the script — not a repeating pattern.
  • Derived from the functional ROLE, not an assumed table name.

Do NOT use:
  • Generic fragments repeated throughout the file.
  • The opening line of any block.
  • Anything that could match at more than one position.

role_anchor_table = the ACTUAL temp table name for that role.

════════════════════════════════════════════════════════════════
MANDATORY JSON RESPONSE FORMAT
════════════════════════════════════════════════════════════════
{DISCOVERY_SCHEMA}

TARGET SQL:
{target_sql}

REMINDER: START WITH {{ END WITH }} — NO TEXT OUTSIDE THE JSON.
""",
        },
    ]
    return [{"role": "user", "content": content_blocks}]


# ─────────────────────────────────────────────────────────────────────────────
# Content-generation prompts
# ─────────────────────────────────────────────────────────────────────────────

PHASE_SCHEMAS = {
    "cleanup_drop": """{
  "insertion_id": <int>,
  "phase": "cleanup_drop",
  "anchor_snippet": "<same anchor from discovery>",
  "insert_position": "after",
  "new_content": "IF OBJECT_ID('tempdb..#bRanked') IS NOT NULL DROP TABLE #bRanked\\nIF OBJECT_ID('tempdb..#bSlots') IS NOT NULL DROP TABLE #bSlots\\nIF OBJECT_ID('tempdb..#LineBreakdown') IS NOT NULL DROP TABLE #LineBreakdown\\nIF OBJECT_ID('tempdb..#recon_temp') IS NOT NULL DROP TABLE #recon_temp"
}""",

    "breakdown_block": """{
  "insertion_id": <int>,
  "phase": "breakdown_block",
  "anchor_snippet": "<same anchor from discovery>",
  "insert_position": "after",
  "confirmed_source_columns": ["col1", "col2", "...every column confirmed present in line_level_source_table"],
  "confirmed_output_columns": ["EncounterID", "Sequence", "ProcedureCode", "...every column in the SELECT INTO #LineBreakdown"],
  "new_content": "<ONLY net-new SQL: SELECT INTO #bRanked, optional SELECT INTO #bSlots, SELECT INTO #LineBreakdown>"
}""",

    "final_join_detail_column": """{
  "insertion_id": <int>,
  "phase": "final_join_detail_column",
  "anchor_snippet": "<same anchor from discovery>",
  "insert_position": "after",
  "new_content": "<ONLY the Price_Breakdown_Detail column — starts with , [Price_Breakdown_Detail] = >"
}""",

    "auditor_queries": """{
  "insertion_id": <int>,
  "phase": "auditor_queries",
  "anchor_snippet": "<same anchor from discovery>",
  "insert_position": "after",
  "new_content": "<ONLY the Phase 4 auditor block: comment banner, Query 1, Query 2 reconciliation>"
}""",
}

PHASE_INSTRUCTIONS = {
    "cleanup_drop": """
Generate THREE DROP TABLE safety lines for the new temp tables introduced by
the breakdown block. Output EXACTLY these three lines in this order:

IF OBJECT_ID('tempdb..#bRanked')       IS NOT NULL DROP TABLE #bRanked
IF OBJECT_ID('tempdb..#bSlots')        IS NOT NULL DROP TABLE #bSlots
IF OBJECT_ID('tempdb..#LineBreakdown') IS NOT NULL DROP TABLE #LineBreakdown
IF OBJECT_ID('tempdb..#recon_temp')    IS NOT NULL DROP TABLE #recon_temp
Nothing else in new_content.""",


    "breakdown_block": """
Generate the full line-level breakdown block using TEMP TABLES (not CTEs).

════════════════════════════════════════════════════════════════
COLUMN SURVIVAL VERIFICATION — MANDATORY BEFORE WRITING ANY SQL
════════════════════════════════════════════════════════════════
Before generating any SQL you MUST:

  1. Identify analysis.line_level_source_table (e.g. #Step3Bundle).
  2. Trace its full construction chain in COMPLETE CURRENT SQL STATE above:
       - Find the SELECT statement that creates it.
       - If it uses SELECT *, trace upstream to find that table's column list.
         Keep tracing through consecutive SELECT * steps until you reach an
         explicit column list. Never assume a column is present without
         tracing all the way back to an explicit list.
       - If it uses an explicit column list, that list is the definitive set.
       - Repeat for every intermediate step back to the source table.
  3. Build a confirmed column list — only columns proven present in the
     final source table may be referenced in the SELECT INTO #LineBreakdown.
  4. For any column your breakdown logic needs that is NOT confirmed present,
     use a NULL placeholder: CAST(NULL AS <appropriate_type>) AS [ColumnName]
     NEVER reference a column that is not confirmed present in the source.
  5. Return confirmed_source_columns and confirmed_output_columns in your JSON.
     These are passed to downstream phases so they know exactly which columns
     exist in #LineBreakdown.

This applies to EVERY column — not just Quantity.

════════════════════════════════════════════════════════════════
GENERATION RULES — USE TEMP TABLES, NOT CTEs
════════════════════════════════════════════════════════════════
Generate THREE sequential temp table blocks:

BLOCK 1 — #bRanked (replaces bRanked CTE):
  SELECT *
      -- Standard MAX category (one winner per encounter):
      , ROW_NUMBER() OVER (PARTITION BY [EncounterID]
            ORDER BY [MAX_CATEGORY] DESC, [Amount] DESC, [Sequence] ASC) AS rn_MAX_ALIAS
      -- HYBRID category (MAX-per-date in AGGREGATE_DATE, SUM in AGGREGATE_ENC):
      --   one winner per encounter per service date — use PARTITION BY EncounterID+ServiceDate.
      --   Do NOT include ServiceDate in the ORDER BY; it is already in the PARTITION.
      , ROW_NUMBER() OVER (PARTITION BY [EncounterID], [ServiceDate]
            ORDER BY [HYBRID_CATEGORY] DESC, [Amount] DESC, [Sequence] ASC) AS rn_HYBRID_ALIAS
      -- WINDOW_REDUCTION: PARTITION BY [EncounterID], ORDER BY SUM of ALL slot columns.
      -- one ROW_NUMBER per MAX category, per HYBRID category, per WINDOW_REDUCTION category.
      -- COMPLETENESS CHECK: count of ROW_NUMBER() lines must equal
      --   (# MAX categories) + (# HYBRID categories) + (# WINDOW_REDUCTION categories).
      --   Verify this count matches the Phase 1C checklist before proceeding.
  INTO #bRanked
  FROM [line_level_source_table]   -- e.g. #Step3Bundle

BLOCK 2 — #bSlots (replaces bSlots CTE — ONLY if analysis.has_window_reduction is true):
  SELECT
      [EncounterID]
      , MAX(CASE WHEN rn_RD  = 1 THEN RD1  ELSE 0 END) AS RD1
      -- ... all slot columns using MAX(CASE WHEN rn_X = 1 THEN SlotCol ELSE 0 END)
  INTO #bSlots
  FROM #bRanked
  GROUP BY [EncounterID]
  -- Use GROUP BY + MAX(CASE...) instead of WHERE rn=1 OR rn=1
  -- This correctly handles multiple window categories in one pass.
  If analysis.has_window_reduction is false, omit #bSlots entirely.

BLOCK 3 — #LineBreakdown (the main output):
  SELECT DISTINCT
      b.[EncounterID], b.[Sequence], b.[ProcedureCode], b.[RevenueCode],
      b.[ServiceDate], b.[Amount] AS [BilledAmount],
      Quantity (NULL placeholder CAST(NULL AS NUMERIC(18,4)) if not confirmed present),
      [ServiceCategory] = CASE ... END   (full hierarchy CASE),
      [PricingMethod]   = CASE ... END   (human-readable rate description),
      [LinePayment]     = ROUND(...)     (MAX/SUM/HYBRID/slot-pivot logic),
      [BundledByNCCI]   = CAST(b.[NeedsCorrectedClaim] AS TINYINT)
                          OR CAST(0 AS TINYINT) if no NCCI bundle
  INTO #LineBreakdown
  FROM #bRanked b
  INNER JOIN [hierarchy_table] s4 ON s4.[EncounterID] = b.[EncounterID]
  LEFT  JOIN #bSlots            rd ON rd.[EncounterID] = b.[EncounterID]
  -- Omit #bSlots join if has_window_reduction is false
  ORDER BY b.[EncounterID], b.[Sequence]

ServiceCategory labeling rules by category type:
  MAX:     IIF(b.rn_CAT = 1, 'CategoryName', 'CategoryName_Non_Winner')
           One winner per encounter. rn>1 = Non_Winner.
  SUM:     'CategoryName' (flat label, all matching lines pay, no rank check)
  HYBRID:  IIF(b.rn_CAT = 1, 'CategoryName', 'CategoryName_Non_Winner')
           HYBRID uses the SAME labeling as MAX — because the PARTITION is
           EncounterID+ServiceDate, rn=1 means winner FOR THAT DATE only.
           Lines with rn>1 on any given date are Non_Winners for that date.
           Multiple rn=1 rows (one per date) all receive the category name.
           DO NOT label all HYBRID lines with the flat category name (SUM
           behavior) — that would obscure which lines actually drove payment.

IMPORTANT NOTES:
  • No CTEs — use temp tables throughout.
  • No semicolon rule needed — temp tables use plain SELECT INTO, no WITH keyword.
  • Do NOT rewrite or repeat any existing SQL.
  • ServiceCategory CASE must cover every category in exact hierarchy order.
  • OP_Default is always last dollar category. INDICATOR_FLAG categories come after.
  • rn_ NAMING CONVENTION — NON-NEGOTIABLE: Every ROW_NUMBER() alias MUST be
    rn_ + FULL category name, NO abbreviations. rn_Chemotherapy not rn_Chemo,
    rn_CardiacCath not rn_CC, rn_GammaKnife not rn_GK. This alias is used
    identically in #bRanked (definition) AND #LineBreakdown (reference). Any
    mismatch or abbreviation causes "Invalid column name" SQL errors at runtime.
  • INDICATOR_FLAG categories (binary 0/1, no dollar value — e.g. COVID, CPT flags):
      - Must have ROW_NUMBER() in #bRanked: PARTITION BY [EncounterID], same as MAX
      - ServiceCategory: IIF(b.rn_FLAG = 1, 'FLAG_Name', 'FLAG_Name_Non_Winner')
        placed AFTER all dollar categories (OP_Default and Suppressed_By_Hierarchy)
      - LinePayment: $0 contribution — DO NOT add flag value to LinePayment sum.
        A binary 1 added to LinePayment creates $1 phantom payments breaking recon.
      - Suppressed_By_Hierarchy sum: MUST include ISNULL(b.[FLAG],0) so pure-flag
        lines get 'Suppressed_By_Hierarchy' not 'No_Payment_Category'.
  • LinePayment rules by category type:
      MAX:              IIF(s4.[CAT] != 0, IIF(b.rn_CAT = 1, ISNULL(b.[CAT], 0), 0), 0)
      SUM:              IIF(s4.[CAT] != 0, ISNULL(b.[CAT], 0), 0)
      HYBRID:           IIF(s4.[CAT] != 0, IIF(b.rn_CAT = 1, ISNULL(b.[CAT], 0), 0), 0)
                        HYBRID uses the SAME rank=1 logic as MAX — because rn is partitioned
                        by EncounterID+ServiceDate, every date's winner (rn=1) contributes
                        its value while non-winners (rn>1) on the same date contribute $0.
                        DO NOT use plain SUM logic for HYBRID — that would pay every line.
      INDICATOR_FLAG:   $0 — no IIF, no contribution to LinePayment sum whatsoever.
      WINDOW_REDUCTION: slot-pivot from #bSlots.

════════════════════════════════════════════════════════════════
CATEGORIES WITH SUB-COLUMNS — CRITICAL SPECIAL HANDLING
════════════════════════════════════════════════════════════════
Some categories are composed of multiple sub-columns at the line level that
get merged into a single column in the encounter-level hierarchy table (#Step4).
You MUST check for this pattern and handle it correctly.

HOW TO DETECT:
  In Phase 1B (LINE_PRICING step), look for any category that has more than
  one column contributing to its total. Common pattern: a main category column
  plus one or more named sub-columns (e.g. [Implants], [Implant_V2790],
  [implant_C1713] are three separate columns in LINE_PRICING that #Step4 merges
  into a single [Implants] column). Check which of these sub-columns survive
  into the line-level source table via the column survival trace.

RULE FOR ServiceCategory:
  When a category has sub-columns, the ServiceCategory WHEN condition must check
  the SUM of all confirmed sub-columns, not just the main column alone:
    WHEN s4.[Implants] != 0
     AND (ISNULL(b.[Implants],0) + ISNULL(b.[Implant_V2790],0) + ISNULL(b.[implant_C1713],0)) > 0
        THEN 'Implants'
  A line with b.[Implants]=0 but b.[Implant_V2790]>0 would incorrectly fall
  through to Suppressed_By_Hierarchy if only b.[Implants] is checked.

RULE FOR LinePayment:
  Similarly, LinePayment must sum all confirmed sub-columns:
    + IIF(s4.[Implants] != 0,
          ISNULL(b.[Implants],0) + ISNULL(b.[Implant_V2790],0) + ISNULL(b.[implant_C1713],0),
          0)
  Using only b.[Implants] will silently produce $0 for V2790/C1713 lines,
  causing reconciliation failures.

DO NOT RE-IMPLEMENT #Step4 SUPPRESSION LOGIC:
  s4.[Implants] already reflects the correct encounter-level suppression outcome
  from #Step4 (including any complex CASE/IIF conditions). If s4.[Implants] != 0,
  the category survived for that encounter — regardless of what other categories
  also survived. Do NOT re-derive suppression by checking s4.[Minor_Surgery],
  s4.[ER], etc. in the ServiceCategory CASE for Implants. Trust s4 as the
  authoritative post-suppression source.""",

    "final_join_detail_column": """Generate ONLY the Price_Breakdown_Detail correlated subquery column per
Phase 3 of the system prompt.

CONFIRMED #LineBreakdown COLUMNS (carry-forward from breakdown_block phase):
{confirmed_lb_columns_block}

CRITICAL: Only reference columns from the confirmed list above when querying
#LineBreakdown. Any column not in that list does not exist in #LineBreakdown
and will cause a SQL error.

  • This is the LAST column in the FINAL_JOIN SELECT.
  • The anchor is the INTO #TableName line itself (e.g. "Into #Step7").
  • insert_position is "before" — the column lands immediately before INTO.
  • Do NOT anchor to the last column definition — anchor to the INTO line.
  • Use STRING_AGG(… , ' || ') WITHIN GROUP (ORDER BY lb.Sequence).
  • Reference #LineBreakdown as lb.
  • Match the encounter alias used in the existing FINAL_JOIN step
    (inspect CURRENT SQL STATE above to find it).
  • Output ONLY the column expression — must start with:
      , [Price_Breakdown_Detail] =
  • Do NOT include surrounding SELECT, INTO, or FROM clauses.""",


    "auditor_queries": """Generate the full Phase 4 auditor block per the system prompt.

CONFIRMED #LineBreakdown COLUMNS (carry-forward from breakdown_block phase):
{confirmed_lb_columns_block}

CRITICAL: All queries may ONLY reference columns from the confirmed list above
when selecting from #LineBreakdown. Any column not in that list does not exist
and will cause a SQL error.

Generate the following four items in order:

  1. Comment banner:
     --======================================================================
     -- *** [AUDITOR PRICE BREAKDOWN SUMMARY QUERIES - INSERTED SECTION] ***
     --======================================================================

  2. Query 1 — Full line-level detail SELECT:
     Standard confirmed columns to include (omit any not in confirmed list):
       lb.[EncounterID], lb.[Sequence], lb.[ProcedureCode], lb.[RevenueCode],
       lb.[ServiceDate], lb.[BilledAmount], lb.[Quantity],
       lb.[ContractYearFrom], lb.[ContractYearTo],
       lb.[ServiceCategory], lb.[PricingMethod],
       lb.[LinePayment] AS [ContractualPayment_Line],
       lb.[BundledByNCCI] AS [ZeroedByNCCI_Flag]

  3. INSERT INTO persistent table — immediately after Query 1 SELECT:
     Insert the same line-level results into [Automation].[dbo].[LineBreakdown_Results].
     Include a ScriptName column derived from the input filename (use the actual
     filename being processed, which can be inferred from the script header comments
     or from the USE [database] statement at the top — use whatever identifier best
     describes this specific contract script).

     INSERT INTO [Automation].[dbo].[LineBreakdown_Results]
     (EncounterID, Sequence, ProcedureCode, RevenueCode, ServiceDate,
      BilledAmount, Quantity, ContractYearFrom, ContractYearTo,
      ServiceCategory, PricingMethod,
      LinePayment, BundledByNCCI, ScriptName, RunDate)
     SELECT
         lb.[EncounterID], lb.[Sequence], lb.[ProcedureCode], lb.[RevenueCode],
         lb.[ServiceDate], lb.[BilledAmount], lb.[Quantity],
         lb.[ContractYearFrom], lb.[ContractYearTo],
         lb.[ServiceCategory], lb.[PricingMethod],
         lb.[LinePayment], lb.[BundledByNCCI],
         '<ScriptIdentifier>',
         GETDATE()
     FROM #LineBreakdown lb

     Only include columns in the INSERT that are confirmed present in #LineBreakdown.
     If Quantity is a NULL placeholder it is still inserted (NULL is valid).

  4. Query 2 — Reconciliation:
     DROP TABLE IF EXISTS #recon_temp
     SELECT EncounterID, SUM(LinePayment) AS LineBreakdown_Total
     INTO #recon_temp FROM #LineBreakdown GROUP BY EncounterID
     Then discrepancy SELECT joining #recon_temp to analysis.total_price_table
     WHERE ABS(LineBreakdown_Total - Price) > 1.00
     ORDER BY ABS discrepancy DESC.

Rules:
  • Use the ACTUAL table names from the analysis for all references.
  • Do NOT repeat or rewrite any existing SQL — purely additive.""",
}




def build_content_messages(system_prompt, reference_sql, current_sql,
                           context, analysis, insertion,
                           confirmed_lb_columns=None):
    """
    confirmed_lb_columns: list of column names confirmed present in #LineBreakdown.
    Populated after breakdown_block phase and passed to subsequent phases so they
    never reference a column that doesn't exist in #LineBreakdown.
    """
    phase        = insertion["phase"]
    anchor       = insertion["anchor_snippet"]
    insertion_id = insertion["insertion_id"]

    schema       = PHASE_SCHEMAS.get(phase, """{
  "insertion_id": <int>,
  "phase": "<phase>",
  "anchor_snippet": "<same anchor>",
  "insert_position": "after",
  "new_content": "<net-new SQL only>"
}""")
    instructions = PHASE_INSTRUCTIONS.get(phase, (
        "Generate ONLY the net-new SQL content needed for this insertion. "
        "Do NOT rewrite or repeat any existing SQL."
    ))

    # Inject confirmed #LineBreakdown columns into downstream phase instructions
    if confirmed_lb_columns and "{confirmed_lb_columns_block}" in instructions:
        col_block = "\n".join(f"  - {c}" for c in confirmed_lb_columns)
        instructions = instructions.replace("{confirmed_lb_columns_block}", col_block)
    elif "{confirmed_lb_columns_block}" in instructions:
        instructions = instructions.replace(
            "{confirmed_lb_columns_block}",
            "  (not yet available — breakdown_block phase has not completed)"
        )

    # Extract retry note if present (set when rn_ consistency check fails)
    retry_note_text = insertion.get("_retry_note", "")
    retry_note = (
        f"\n⚠️ RETRY INSTRUCTION — READ THIS FIRST:\n{retry_note_text}\n\n"
        if retry_note_text else ""
    )

    content_blocks = [
        {
            "type": "text",
            "text": system_prompt,
            "cache_control": {"type": "ephemeral"},
        },
        {
            "type": "text",
            "text": f"REFERENCE SQL:\n{reference_sql}",
            "cache_control": {"type": "ephemeral"},
        },
        {
            "type": "text",
            "text": f"""CRITICAL: RETURN ONLY VALID JSON — NO PROSE, NO MARKDOWN, NO EXPLANATION.

════════════════════════════════════════════════════════════════
SCRIPT ANALYSIS  (from discovery — treat as authoritative)
════════════════════════════════════════════════════════════════
{json.dumps(analysis, indent=2)}

════════════════════════════════════════════════════════════════
COMPLETE CURRENT SQL STATE
(verify column names, table aliases, existing structure here)
════════════════════════════════════════════════════════════════
{current_sql}

════════════════════════════════════════════════════════════════
INSERTION TASK  —  insertion_id={insertion_id}  phase={phase}
════════════════════════════════════════════════════════════════
{retry_note}{instructions}

LOCAL CONTEXT AROUND INSERTION POINT:
{context}

ANCHOR (insertion happens immediately after this):
{anchor}

════════════════════════════════════════════════════════════════
MANDATORY JSON RESPONSE FORMAT
════════════════════════════════════════════════════════════════
{schema}

REMINDER: new_content = NET-NEW SQL ONLY.  START WITH {{ END WITH }}.
""",
        },
    ]
    return [{"role": "user", "content": content_blocks}]


# ─────────────────────────────────────────────────────────────────────────────
# Claude API call with retry
# ─────────────────────────────────────────────────────────────────────────────

def call_claude_with_retry(client, messages, label, max_retries=3, max_tokens=40000):
    global total_cached_tokens, total_uncached_tokens

    for attempt in range(max_retries + 1):
        try:
            print(f"\n  ── {label} (attempt {attempt + 1}) ──")
            t0 = time.perf_counter()

            text = ""
            msg = None
            
            try:
                msg  = client.messages.create(
                    model=MODEL, max_tokens=max_tokens, messages=messages
                )
                text = "".join(
                    b.text for b in msg.content
                    if getattr(b, "type", None) == "text"
                )
                print(f"  Standard response received ({len(text)} chars)")
                
            except Exception as e:
                if "Streaming is required" not in str(e):
                    raise
                    
                print("  Switching to streaming…")
                text = ""
                stream_chunks = 0
                
                try:
                    with client.messages.stream(
                        model=MODEL, max_tokens=max_tokens, messages=messages
                    ) as stream:
                        for chunk in stream:
                            if (chunk.type == "content_block_delta"
                                    and hasattr(chunk.delta, "text")):
                                text += chunk.delta.text
                                stream_chunks += 1
                                
                                # Progress indicator for large streams
                                if stream_chunks % 100 == 0:
                                    print(f"    Streaming... {len(text)} chars received")
                        
                        msg = stream.get_final_message()
                        print(f"  Streaming complete ({stream_chunks} chunks, {len(text)} chars)")
                        
                except Exception as stream_error:
                    print(f"  Streaming error: {stream_error}")
                    raise

            # Only attempt JSON parsing after we have the complete response
            if not text.strip():
                raise ValueError("Empty response received from Claude")

            elapsed = time.perf_counter() - t0
            
            # Parse JSON from complete response
            parsed = extract_json(text)

            # Check if parsing failed and we should retry
            if "error" in parsed and attempt < max_retries:
                error_msg = parsed.get("error", "Unknown JSON error")
                print(f"  JSON parse failed: {error_msg}")
                print(f"  Response length: {len(text)} chars")
                print(f"  Response preview: {text[:200]}...")
                print(f"  Retrying in 1.5s…")
                time.sleep(1.5)
                continue

            # Calculate usage and cost
            inp, out, _, cache_read = usage_numbers(msg) if msg else (0, 0, 0, 0)
            total_cached_tokens   += cache_read
            total_uncached_tokens += inp
            cost = estimate_cost(inp, out, cache_read)

            print(f"  in={inp:,} out={out:,} cached={cache_read:,} "
                  f"cost=${cost:.4f} time={elapsed:.1f}s")
            
            return parsed, msg, elapsed, cost

        except Exception as e:
            err = str(e)
            print(f"  API error: {err}")
            if attempt < max_retries:
                wait = (2 ** attempt + random.uniform(0, 1)) if "429" in err else 1.5
                print(f"  Retrying in {wait:.1f}s…")
                time.sleep(wait)
            else:
                raise


# ─────────────────────────────────────────────────────────────────────────────
# Required-phase validation
# ─────────────────────────────────────────────────────────────────────────────

def check_required_phases(insertions: list, log_path) -> list[str]:
    found   = {ins.get("phase") for ins in insertions}
    missing = sorted(REQUIRED_PHASES - found)
    if missing:
        msg = f"Missing required phases: {missing}"
        print(f"\n  WARNING: {msg}")
        append_log(log_path, f"MISSING_PHASES | {msg}")
    return missing


# ─────────────────────────────────────────────────────────────────────────────
# Single-file processor
# ─────────────────────────────────────────────────────────────────────────────

def process_single_file(client, input_file_path, system_prompt,
                        reference_sql, file_outputs, session_id):
    global db_manager

    print(f"\n{'='*70}")
    print(f"  Processing: {input_file_path.name}")
    print(f"{'='*70}")

    file_id    = None
    file_start = time.perf_counter()

    if db_manager:
        try:
            file_id = db_manager.start_file(
                session_id=session_id,
                filename=input_file_path.name,
                input_file_path=str(input_file_path),
                file_size_bytes=input_file_path.stat().st_size,
            )
        except Exception as e:
            print(f"  DB error starting file: {e}")

    insertions_log = {
        "session_info": {
            "timestamp":  datetime.now().isoformat(),
            "input_file": str(input_file_path),
            "approach":   "Role-based anchor insertion v9 — INDICATOR_FLAG + rn_ naming + consistency retry",
            "file_id":    file_id,
            "session_id": session_id,
        },
        "analysis":       {},
        "missing_phases": [],
        "insertions":     [],
    }

    file_total_input  = 0
    file_total_output = 0
    file_total_time   = 0.0
    file_total_cost   = 0.0

    try:
        target_sql = read_text(input_file_path)

        # ── Phase 1: Discovery + analysis ────────────────────────────────────
        print("\n  Phase 1 — Discovery & role analysis…")
        discovery_msgs = build_discovery_messages(system_prompt, reference_sql, target_sql)
        discovery, msg, elapsed, cost = call_claude_with_retry(
            client, discovery_msgs, "DISCOVERY"
        )

        inp, out, _, _ = usage_numbers(msg)
        file_total_input  += inp
        file_total_output += out
        file_total_time   += elapsed
        file_total_cost   += cost

        if "error" in discovery:
            raise RuntimeError(f"Discovery JSON parse failed: {discovery['error']}")

        analysis   = discovery.get("analysis", {})
        insertions = discovery.get("insertions", [])

        # Check required phases
        missing_phases = check_required_phases(insertions, file_outputs["log_file"])
        insertions_log["missing_phases"] = missing_phases

        # Save analysis + plan for audit
        insertions_log["analysis"] = analysis
        with open(file_outputs["analysis_log"], "w", encoding="utf-8") as f:
            json.dump({
                "analysis":       analysis,
                "insertion_plan": insertions,
                "missing_phases": missing_phases,
            }, f, indent=2)

        print(f"\n  Analysis:")
        print(f"    line_level_source  : {analysis.get('line_level_source_table', '?')}")
        print(f"    hierarchy_table    : {analysis.get('hierarchy_table', '?')}")
        print(f"    total_price_table  : {analysis.get('total_price_table', '?')}")
        print(f"    final_join_table   : {analysis.get('final_join_table', '?')}")
        print(f"    output_table       : {analysis.get('output_table', '?')}")
        print(f"    has_ncci_bundle    : {analysis.get('has_ncci_bundle', '?')}")
        print(f"    has_window_reduct  : {analysis.get('has_window_reduction', '?')}")
        print(f"    has_agg_date_step  : {analysis.get('has_aggregate_date_step', '?')}")
        print(f"    max_categories     : {analysis.get('max_categories', [])}")
        print(f"    hybrid_categories  : {analysis.get('hybrid_categories', [])}")
        print(f"    indicator_flags    : {analysis.get('indicator_flag_categories', [])}")
        print(f"    semicolon_needed   : {analysis.get('semicolon_needed_before_cte', '?')}")
        print(f"    insertions planned : {len(insertions)}")
        if missing_phases:
            print(f"    ⚠  MISSING phases  : {missing_phases}")

        if not insertions:
            print("  No insertions planned — saving original.")
            file_outputs["patched_sql"].write_text(target_sql, encoding="utf-8")
            _save_log(file_outputs, insertions_log, input_file_path,
                      0, 0, 0, file_total_input, file_total_output,
                      file_total_time, file_total_cost)
            _end_db_file(file_id, "SUCCESS", file_outputs, 0, 0,
                         file_total_cost, time.perf_counter() - file_start)
            return _result(input_file_path.name, True, 0, 0,
                           file_total_cost, time.perf_counter() - file_start)

        # ── Phase 2: Generate + validate + apply ─────────────────────────────
        print(f"\n  Phase 2 — Generating {len(insertions)} insertion(s)…")

        current_sql = target_sql
        successful  = 0
        failed      = 0

        confirmed_lb_columns = []  # populated after breakdown_block, passed to downstream phases
        for i, ins in enumerate(insertions):
            attempt_count = 0  # tracks retries for breakdown_block rn_ consistency
            insertion_id = ins.get("insertion_id", i + 1)
            phase        = ins.get("phase", "unknown")
            description  = ins.get("description", "")
            anchor       = ins.get("anchor_snippet", "")
            position     = ins.get("insert_position", "after")
            role_table   = ins.get("role_anchor_table", "")

            print(f"\n  [{i+1}/{len(insertions)}] phase={phase}")
            print(f"  {description}")

            # ── Anchor validation ─────────────────────────────────────────────
            if not anchor:
                _skip(file_id, insertion_id, description, anchor, position,
                      "Empty anchor_snippet", file_outputs)
                failed += 1
                continue

            anchor_pos = find_anchor(current_sql, anchor)
            if anchor_pos is None:
                _skip(file_id, insertion_id, description, anchor, position,
                      "Anchor not found in current SQL", file_outputs,
                      preview=anchor[:100])
                failed += 1
                continue

            # Compute anchor line number BEFORE insertion
            anchor_line = get_line_number(current_sql, anchor_pos)
            print(f"  Anchor found at line {anchor_line}")

            # ── Idempotency guard ─────────────────────────────────────────────
            PHASE_SIGNATURES = {
                "cleanup_drop":             "IF OBJECT_ID('tempdb..#LineBreakdown')",
                "breakdown_block":          "INTO #LineBreakdown",
                "final_join_detail_column": "[Price_Breakdown_Detail]",
                "auditor_queries":          "AUDITOR PRICE BREAKDOWN SUMMARY QUERIES",
            }
            signature = PHASE_SIGNATURES.get(phase)
            if signature and signature.upper() in current_sql.upper():
                print(f"  SKIP (idempotent): phase={phase} signature already present in SQL")
                append_log(file_outputs["log_file"], f"IDEMPOTENT_SKIP | {phase} | signature found")
                continue

            # Role-table proximity check (non-fatal warning)
            if role_table and role_table != "cleanup_block":
                if not validate_anchor_against_table(current_sql, anchor, role_table):
                    warn = (f"role_table '{role_table}' not in 3000-char window "
                            f"before anchor — possible drift")
                    print(f"  WARNING: {warn}")
                    append_log(file_outputs["log_file"],
                               f"ANCHOR_WARNING | {phase} | {warn}")

            # ── Generate content ──────────────────────────────────────────────
            context      = get_context(current_sql, anchor)
            content_msgs = build_content_messages(
                system_prompt, reference_sql, current_sql,
                context, analysis, ins,
                confirmed_lb_columns=confirmed_lb_columns,
            )

            ins_start = time.perf_counter()
            try:
                # Per-phase token limits.
                # breakdown_block generates #bRanked + #bSlots + full #LineBreakdown
                # SELECT with ServiceCategory/PricingMethod/LinePayment CASEs —
                # easily 600-800 lines, which saturates 16k tokens.
                # Other phases produce small outputs and 16k is sufficient.
                PHASE_MAX_TOKENS = {
                    "breakdown_block":          40000,
                    "auditor_queries":          8000,
                    "final_join_detail_column": 4000,
                    "cleanup_drop":             1000,
                }
                phase_max_tokens = PHASE_MAX_TOKENS.get(phase, 40000)

                rewrite, msg, elapsed, cost = call_claude_with_retry(
                    client, content_msgs, f"GENERATE {phase}",
                    max_tokens=phase_max_tokens,
                )

                inp, out, _, _ = usage_numbers(msg)
                file_total_input  += inp
                file_total_output += out
                file_total_time   += elapsed
                file_total_cost   += cost

                new_content = rewrite.get("new_content", "")
                if not new_content or not new_content.strip():
                    _skip(file_id, insertion_id, description, anchor, position,
                          "Empty new_content returned", file_outputs,
                          anchor_line=anchor_line)
                    failed += 1
                    continue

                # ── Post-generation validation for breakdown_block ────────
                if phase == "breakdown_block":
                    missing_rn = validate_branked_consistency(new_content)
                    if missing_rn:
                        print(f"  ⚠ rn_ consistency check FAILED — referenced but not defined: {missing_rn}")
                        append_log(file_outputs["log_file"],
                                   f"RN_CONSISTENCY_FAIL | {missing_rn}")
                        if attempt_count < 2:  # retry once
                            attempt_count += 1
                            print(f"  Retrying breakdown_block generation (attempt {attempt_count+1})…")
                            # Add the missing aliases to the instruction context
                            ins_retry = dict(ins)
                            ins_retry["_retry_note"] = (
                                f"CRITICAL RETRY: The previous attempt defined these rn_ aliases "
                                f"in #bRanked but FAILED to include: {missing_rn}. "
                                f"You MUST add a ROW_NUMBER() line for each missing alias "
                                f"in the #bRanked SELECT before INTO #bRanked. "
                                f"Every rn_ alias referenced in #LineBreakdown MUST be defined in #bRanked."
                            )
                            content_msgs_retry = build_content_messages(
                                system_prompt, reference_sql, current_sql,
                                context, analysis, ins_retry,
                                confirmed_lb_columns=confirmed_lb_columns,
                            )
                            rewrite, msg, elapsed2, cost2 = call_claude_with_retry(
                                client, content_msgs_retry, f"RETRY {phase}",
                                max_tokens=phase_max_tokens,
                            )
                            inp2, out2, _, _ = usage_numbers(msg)
                            file_total_input  += inp2
                            file_total_output += out2
                            file_total_time   += elapsed2
                            file_total_cost   += cost2
                            new_content = rewrite.get("new_content", "")
                            # Check again
                            missing_rn2 = validate_branked_consistency(new_content)
                            if missing_rn2:
                                print(f"  ⚠ rn_ still missing after retry: {missing_rn2}")
                                append_log(file_outputs["log_file"],
                                           f"RN_CONSISTENCY_FAIL_RETRY | {missing_rn2}")
                            else:
                                print(f"  ✓ rn_ consistency check passed after retry")

                # ── Apply ─────────────────────────────────────────────────────
                current_sql, success = apply_insertion(
                    current_sql, anchor, new_content, position
                )
                lines_added         = len(new_content.strip().splitlines())
                chars_added         = len(new_content)
                insertion_ends_line = anchor_line + lines_added if success else None

                if success:
                    successful += 1
                    print(f"  ✓ Inserted {lines_added} lines / {chars_added} chars")
                    print(f"    Lines {anchor_line} → {insertion_ends_line} in patched SQL")

                    # Extract confirmed #LineBreakdown columns after breakdown_block
                    # so downstream phases (final_join_detail_column, auditor_queries)
                    # know exactly which columns exist and can reference safely.
                    if phase == "breakdown_block":
                        confirmed_lb_columns = rewrite.get("confirmed_output_columns", [])
                        if confirmed_lb_columns:
                            print(f"  ✓ Confirmed #LineBreakdown columns: {confirmed_lb_columns}")
                            append_log(file_outputs["log_file"],
                                       f"LB_COLUMNS | {confirmed_lb_columns}")
                        else:
                            print("  WARNING: breakdown_block returned no confirmed_output_columns")
                            append_log(file_outputs["log_file"],
                                       "LB_COLUMNS_MISSING | breakdown_block did not return confirmed_output_columns")

                    insertions_log["insertions"].append({
                        "insertion_id":          insertion_id,
                        "phase":                 phase,
                        "description":           description,
                        "role_table":            role_table,
                        "anchor_preview":        anchor[:150],
                        "anchor_line_number":    anchor_line,
                        "insertion_ends_at_line": insertion_ends_line,
                        "lines_added":           lines_added,
                        "chars_added":           chars_added,
                        "new_content_preview":   new_content[:600],
                        **({"confirmed_output_columns": confirmed_lb_columns}
                           if phase == "breakdown_block" else {}),
                    })
                else:
                    failed += 1
                    print(f"  ✗ apply_insertion returned False for {phase}")

                _log_db_insertion(
                    file_id, insertion_id, description,
                    anchor, position,
                    anchor_line, insertion_ends_line,
                    success, new_content,
                    lines_added if success else 0,
                    chars_added if success else 0,
                    elapsed,
                    None if success else "apply_insertion returned False",
                )

                append_log(
                    file_outputs["log_file"],
                    f"INSERTION {insertion_id} | phase={phase} "
                    f"anchor_line={anchor_line} ends_line={insertion_ends_line} "
                    f"in={inp} out={out} time={elapsed:.2f} "
                    f"cost={cost:.4f} success={success} lines={lines_added}",
                )

            except Exception as e:
                print(f"  ERROR: {e}")
                failed += 1
                _log_db_insertion(
                    file_id, insertion_id, description,
                    anchor, position,
                    anchor_line, None,
                    False, "", 0, 0,
                    time.perf_counter() - ins_start, str(e),
                )

        # ── Save outputs ──────────────────────────────────────────────────────
        file_outputs["patched_sql"].write_text(current_sql, encoding="utf-8")
        _save_log(file_outputs, insertions_log, input_file_path,
                  len(insertions), successful, failed,
                  file_total_input, file_total_output,
                  file_total_time, file_total_cost)

        file_runtime = time.perf_counter() - file_start

        if failed == 0 and not missing_phases:
            status = "SUCCESS"
        elif successful > 0:
            status = "PARTIAL"
        else:
            status = "FAILED"

        _end_db_file(file_id, status, file_outputs,
                     len(insertions), successful, file_total_cost, file_runtime)

        print(f"\n  Result   : {successful}/{len(insertions)} successful | "
              f"${file_total_cost:.4f} | {file_runtime:.1f}s")
        if missing_phases:
            print(f"  ⚠ Missing: {missing_phases}")
        print(f"  Output   : {file_outputs['patched_sql']}")
        print(f"  Analysis : {file_outputs['analysis_log']}")

        return _result(input_file_path.name, status == "SUCCESS",
                       len(insertions), successful,
                       file_total_cost, file_runtime,
                       missing_phases=missing_phases)

    except Exception as e:
        print(f"  FATAL: {e}")
        append_log(file_outputs["log_file"], f"ERROR | {input_file_path.name} | {e}")
        file_runtime = time.perf_counter() - file_start
        _end_db_file(file_id, "FAILED", file_outputs,
                     0, 0, 0.0, file_runtime, str(e))
        return _result(input_file_path.name, False, 0, 0,
                       0.0, file_runtime, error=str(e))


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _skip(file_id, ins_id, desc, anchor, position, reason, file_outputs,
          preview=None, anchor_line=None):
    print(f"  SKIP: {reason}")
    if preview:
        print(f"  Anchor preview: {preview}")
    append_log(file_outputs["log_file"],
               f"SKIP | insertion={ins_id} | {reason}")
    _log_db_insertion(file_id, ins_id, desc, anchor, position,
                      anchor_line, None,
                      False, "", 0, 0, 0.0, reason)


def _save_log(file_outputs, insertions_log, input_file_path,
              total, successful, failed, inp, out, t, cost):
    with open(file_outputs["insertions_log"], "w", encoding="utf-8") as f:
        json.dump(insertions_log, f, indent=2)
    append_log(
        file_outputs["log_file"],
        f"FILE_SUMMARY | {input_file_path.name} | "
        f"insertions={total} success={successful} failed={failed} "
        f"input={inp} output={out} time={t:.2f} cost={cost:.4f}",
    )


def _end_db_file(file_id, status, file_outputs, found, applied,
                 cost, secs, error=None):
    global db_manager
    if db_manager and file_id:
        try:
            db_manager.end_file(
                file_id=file_id,
                status=status,
                output_file_path=str(file_outputs["patched_sql"]),
                edits_found=found,
                edits_applied=applied,
                cost=cost,
                processing_seconds=secs,
                **({"error_message": error} if error else {}),
            )
        except Exception as e:
            print(f"  DB error ending file: {e}")


def _log_db_insertion(file_id, ins_id, desc,
                      anchor, position,
                      anchor_line, insertion_ends_line,
                      success, new_content,
                      lines, chars, secs, error):
    """
    Log insertion to database.

    DB schema required:
      position               NVARCHAR(500)  — anchor text (not integer offset)
      anchor_line_number     INT NULL       — line where anchor ends in source SQL
      insertion_ends_at_line INT NULL       — anchor_line + lines_added in patched SQL

    Run once before first use:
      ALTER TABLE dbo.SQL_Processing_Edits
          ALTER COLUMN position NVARCHAR(500);
      ALTER TABLE dbo.SQL_Processing_Edits
          ADD anchor_line_number    INT NULL,
              insertion_ends_at_line INT NULL;
    """
    global db_manager
    if db_manager and file_id:
        try:
            # Truncate description to fit database column constraints
            # Most SQL databases have limited varchar sizes for description fields
            max_desc_length = 255  # Conservative limit for description column
            truncated_desc = desc[:max_desc_length] if desc else ""
            if len(desc) > max_desc_length:
                truncated_desc = desc[:max_desc_length-3] + "..."
                print(f"  Description truncated: {len(desc)} -> {len(truncated_desc)} chars")

            db_manager.log_edit(
                file_id=file_id,
                edit_number=ins_id,
                description=truncated_desc,          # Truncated to fit column
                position=anchor[:500],               # NVARCHAR(500) — anchor text
                match_type=position,
                success=success,
                old_snippet_full_text=anchor,
                new_snippet_full_text=new_content,
                lines_added=lines,
                characters_added=chars,
                processing_seconds=secs,
                error_message=error,
                anchor_line_number=anchor_line,
                insertion_ends_at_line=insertion_ends_line,
            )
        except Exception as e:
            print(f"  DB error logging insertion: {e}")


def _result(filename, success, found, applied, cost, t,
            error=None, missing_phases=None):
    r = {
        "filename":           filename,
        "success":            success,
        "insertions_found":   found,
        "insertions_applied": applied,
        "cost":               cost,
        "time":               t,
        "missing_phases":     missing_phases or [],
    }
    if error:
        r["error"] = error
    return r


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    global session_metrics, db_manager, total_cached_tokens, total_uncached_tokens

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("Set ANTHROPIC_API_KEY in .env")

    client = anthropic.Anthropic(api_key=api_key)
    session_metrics["start_time"] = time.perf_counter()

    print("Claude SQL Batch INSERT Pipeline  v9")
    print("Strategy: INDICATOR_FLAG support, rn_ naming enforcement, rn_ consistency retry")

    try:
        db_manager = DatabaseManager()
        ok = db_manager.test_connection()
        print(f"Database: {'OK' if ok else 'FAILED — no DB logging'}")
        if not ok:
            db_manager = None
    except Exception as e:
        print(f"Database unavailable ({e}) — no DB logging")
        db_manager = None

    INPUT_DIR.mkdir(exist_ok=True)
    BATCH_OUTPUT_DIR.mkdir(exist_ok=True)

    sql_files = list(INPUT_DIR.glob("*.sql"))
    if not sql_files:
        print(f"No .sql files in {INPUT_DIR}")
        return

    print(f"\nFiles to process: {len(sql_files)}")
    for f in sql_files:
        print(f"  {f.name}")

    print("\nLoading system prompt + reference SQL…")
    system_prompt = read_docx(SYSTEM_PROMPT_PATH)
    reference_sql = read_text(REFERENCE_SQL_PATH)

    session_id = None
    if db_manager:
        try:
            session_id = db_manager.start_session(
                script_version="sql_insert_pipeline_v9.py",
                reference_file=REFERENCE_SQL_PATH.name,
                approach="Role-based anchor insertion v9 — INDICATOR_FLAG + rn_ naming + consistency retry",
            )
            print(f"DB session: {session_id}")
        except Exception as e:
            print(f"DB session start error: {e}")

    results = []
    for sql_file in sql_files:
        file_outputs = create_file_output_structure(sql_file.name, BATCH_OUTPUT_DIR)
        result       = process_single_file(
            client, sql_file, system_prompt, reference_sql,
            file_outputs, session_id,
        )
        results.append(result)

        session_metrics["files_processed"] += 1
        if result["success"]:
            session_metrics["files_successful"]            += 1
            session_metrics["total_insertions"]            += result["insertions_found"]
            session_metrics["total_successful_insertions"] += result["insertions_applied"]
        else:
            session_metrics["files_failed"] += 1
        session_metrics["total_cost"] += result["cost"]

    total_time      = time.perf_counter() - session_metrics["start_time"]
    all_tokens      = total_uncached_tokens + total_cached_tokens
    cost_no_cache   = (all_tokens / 1_000_000) * INPUT_COST_PER_MILLION
    cost_with_cache = session_metrics["total_cost"]
    savings         = cost_no_cache - cost_with_cache
    savings_pct     = (savings / cost_no_cache * 100) if cost_no_cache > 0 else 0

    if db_manager and session_id:
        try:
            db_manager.end_session(
                session_id=session_id,
                status="COMPLETED" if session_metrics["files_failed"] == 0 else "PARTIAL",
                total_files=session_metrics["files_processed"],
                files_successful=session_metrics["files_successful"],
                files_failed=session_metrics["files_failed"],
                total_edits=session_metrics["total_insertions"],
                total_successful_edits=session_metrics["total_successful_insertions"],
                total_cost=cost_with_cache,
                total_savings=savings,
                processing_seconds=total_time,
            )
            db_manager.disconnect()
        except Exception as e:
            print(f"DB session end error: {e}")

    print(f"\n{'='*60}")
    print("BATCH SUMMARY")
    print(f"  Files      : {session_metrics['files_successful']}/{session_metrics['files_processed']} successful")
    print(f"  Insertions : {session_metrics['total_successful_insertions']}/{session_metrics['total_insertions']} applied")
    print(f"  Cost       : ${cost_with_cache:.4f}  (saved ${savings:.4f} / {savings_pct:.0f}% via cache)")
    print(f"  Time       : {total_time:.1f}s")
    print()
    for r in results:
        icon = "✓" if r["success"] else "✗"
        missing_note = f" ⚠ missing={r['missing_phases']}" if r.get("missing_phases") else ""
        print(f"  {icon} {r['filename']} | "
              f"{r['insertions_applied']}/{r['insertions_found']} | "
              f"${r['cost']:.4f} | {r['time']:.1f}s{missing_note}")
        if not r["success"] and r.get("error"):
            print(f"    Error: {r['error']}")

    summary_log = BATCH_OUTPUT_DIR / "batch_summary.log"
    with open(summary_log, "w", encoding="utf-8") as f:
        f.write(
            f"session_id={session_id} "
            f"files={session_metrics['files_processed']} "
            f"success={session_metrics['files_successful']} "
            f"failed={session_metrics['files_failed']} "
            f"insertions={session_metrics['total_insertions']} "
            f"applied={session_metrics['total_successful_insertions']} "
            f"cost={cost_with_cache:.4f} savings={savings:.4f} "
            f"time={total_time:.2f}\n\nFILE RESULTS:\n"
        )
        for r in results:
            f.write(
                f"{r['filename']} | success={r['success']} | "
                f"insertions={r['insertions_applied']}/{r['insertions_found']} | "
                f"missing={r['missing_phases']} | "
                f"cost=${r['cost']:.4f}\n"
            )

    print(f"\nOutputs → {BATCH_OUTPUT_DIR}")


if __name__ == "__main__":
    main()